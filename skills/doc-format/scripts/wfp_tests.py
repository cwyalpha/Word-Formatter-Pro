# -*- coding: utf-8 -*-
"""Built-in unit tests for Word Formatter Pro."""

from __future__ import annotations

import os
import subprocess
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from wfp_config import DEFAULT_CONFIG
from wfp_core import (
    BLANK_LINE_MODE_DELETE_SINGLE,
    BLANK_LINE_MODE_KEEP_SINGLE,
    BLANK_LINE_MODE_PRESERVE,
    LegacyConversionUnavailable,
    SofficeConverter,
    WordProcessor,
)


class TextNormalizationTests(unittest.TestCase):
    def test_symbol_normalization_keeps_decimal_numbers(self):
        self.assertEqual(
            WordProcessor._normalize_symbols_in_text("你好,世界."),
            "你好，世界。",
        )
        self.assertEqual(
            WordProcessor._normalize_symbols_in_text("3.14 是 pi"),
            "3.14 是 pi",
        )

    def test_ellipsis_and_quotes(self):
        self.assertEqual(
            WordProcessor._normalize_symbols_in_text('他说"你好"...'),
            "他说“你好”……",
        )
        self.assertEqual(
            WordProcessor._normalize_symbols_in_text("version 1.2"),
            "version 1.2",
        )

    def test_markdown_cleaning(self):
        raw = "# 标题\n**粗体** 和 [链接](https://example.com)\n![图片](a.png)\n> 引用"
        cleaned = WordProcessor._clean_markdown(raw)
        self.assertEqual(cleaned, "标题\n粗体 和 链接\n图片\n引用")

    def test_markdown_cleaning_preserves_source_numeric_numbering(self):
        raw = (
            "一、登录\n\n"
            "1. 打开软件\n"
            "2. 完成登录\n\n"
            "（一）首页搜索\n\n"
            "1. 输入关键词\n"
            "2. 点击院校\n"
            "1.2.3 小节号保持原样"
        )
        cleaned = WordProcessor._clean_markdown(raw)
        self.assertEqual(cleaned, raw)


class BlankLineTests(unittest.TestCase):
    def test_delete_single_blank_line_and_compress_multiple(self):
        text = "a\n\nb\n\n\nc"
        self.assertEqual(
            WordProcessor._remove_blank_lines_from_text(text),
            "a\nb\n\nc",
        )

    def test_keep_single_blank_line_and_compress_multiple(self):
        text = "a\n\nb\n\n\nc"
        self.assertEqual(
            WordProcessor._remove_blank_lines_from_text(text, keep_single_blank_lines=True),
            "a\n\nb\n\nc",
        )

    def test_blank_line_mode_aliases(self):
        self.assertEqual(
            WordProcessor._normalize_blank_line_mode("preserve"),
            BLANK_LINE_MODE_PRESERVE,
        )
        self.assertEqual(
            WordProcessor._normalize_blank_line_mode("delete_single"),
            BLANK_LINE_MODE_DELETE_SINGLE,
        )
        self.assertEqual(
            WordProcessor._normalize_blank_line_mode("keep_single"),
            BLANK_LINE_MODE_KEEP_SINGLE,
        )

    def test_processor_preserve_mode_leaves_text_unchanged(self):
        processor = WordProcessor({"blank_line_mode": BLANK_LINE_MODE_PRESERVE})
        text = "a\n\nb\n\n\nc"
        self.assertEqual(processor._normalize_text_blank_lines(text), text)


class TableHelperTests(unittest.TestCase):
    def test_numeric_table_text(self):
        self.assertTrue(WordProcessor._is_numeric_table_text("1,234.56"))
        self.assertTrue(WordProcessor._is_numeric_table_text("¥100元"))
        self.assertTrue(WordProcessor._is_numeric_table_text("12.5%"))
        self.assertFalse(WordProcessor._is_numeric_table_text("abc"))

    def test_short_table_text(self):
        self.assertTrue(WordProcessor._is_short_table_text("合计", max_len=4))
        self.assertFalse(WordProcessor._is_short_table_text("较长文本内容", max_len=4))

    def test_table_percentage_normalization(self):
        pcts = WordProcessor._normalize_table_pcts([1, 9], 20, 80)
        self.assertAlmostEqual(sum(pcts), 100.0)
        self.assertEqual(pcts, [20.0, 80.0])


class OoxmlProtectionTests(unittest.TestCase):
    def test_ooxml_element_detection(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run()
        run._r.append(OxmlElement("w:drawing"))
        self.assertTrue(WordProcessor._has_drawing_or_pict(para))
        self.assertFalse(WordProcessor._has_field_codes(para))

        para_field = doc.add_paragraph()
        field_run = para_field.add_run()
        field_run._r.append(OxmlElement("w:fldChar"))
        self.assertTrue(WordProcessor._has_field_codes(para_field))

    def test_strip_leading_whitespace_removes_plain_blank_run(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("   ")
        para.add_run("正文")
        WordProcessor({})._strip_leading_whitespace(para)
        self.assertEqual(para.text, "正文")

    def test_strip_leading_whitespace_preserves_special_run(self):
        doc = Document()
        para = doc.add_paragraph()
        special_run = para.add_run()
        special_run._r.append(OxmlElement("w:fldChar"))
        para.add_run("正文")

        WordProcessor({})._strip_leading_whitespace(para)

        self.assertTrue(WordProcessor._has_field_codes(para))
        self.assertEqual(len(para.runs), 2)
        self.assertEqual(para.text, "正文")


class TempAndConversionTests(unittest.TestCase):
    def test_temp_docx_path_uses_system_temp_and_safe_name(self):
        processor = WordProcessor(DEFAULT_CONFIG.copy())
        temp_path = processor._make_temp_docx_path("copy", 'bad<name>:"?')
        try:
            self.assertEqual(
                os.path.normcase(os.path.abspath(os.path.dirname(temp_path))),
                os.path.normcase(os.path.abspath(tempfile.gettempdir())),
            )
            self.assertRegex(os.path.basename(temp_path), r"^~temp_copy_bad_name_[0-9]+_[0-9a-f]{8}\.docx$")
            self.assertIn(temp_path, processor.temp_files)
        finally:
            processor._cleanup_temp_files()

    def test_txt_conversion_uses_blank_line_mode_and_cleans_temp_file(self):
        with tempfile.TemporaryDirectory(prefix="wfp273_test_") as tmpdir:
            source = Path(tmpdir) / "sample.txt"
            source.write_text("标题\n\n正文一\n\n\n正文二", encoding="utf-8")

            processor = WordProcessor(DEFAULT_CONFIG.copy())
            temp_docx, is_from_txt = processor.convert_to_docx(str(source))
            try:
                self.assertTrue(is_from_txt)
                self.assertTrue(os.path.exists(temp_docx))
                self.assertEqual(
                    os.path.normcase(os.path.abspath(os.path.dirname(temp_docx))),
                    os.path.normcase(os.path.abspath(tempfile.gettempdir())),
                )
                self.assertFalse(list(Path(tmpdir).glob("~temp_*.docx")))

                converted = Document(temp_docx)
                self.assertEqual(
                    [p.text for p in converted.paragraphs],
                    ["标题", "正文一", "", "正文二"],
                )
            finally:
                processor._cleanup_temp_files()
            self.assertFalse(os.path.exists(temp_docx))

    def test_missing_soffice_marks_legacy_conversion_skipped(self):
        class MissingConverter:
            available = False

        processor = WordProcessor(DEFAULT_CONFIG.copy())
        processor.soffice_converter = MissingConverter()
        with self.assertRaises(LegacyConversionUnavailable):
            processor._convert_legacy_with_soffice("legacy.doc", "unused.docx")

    def test_format_document_skips_wps_when_soffice_missing(self):
        class MissingConverter:
            available = False

        with tempfile.TemporaryDirectory(prefix="wfp_legacy_skip_test_") as tmpdir:
            source = Path(tmpdir) / "legacy.wps"
            source.write_bytes(b"not a real wps file")

            processor = WordProcessor(DEFAULT_CONFIG.copy())
            processor.soffice_converter = MissingConverter()
            with self.assertRaises(LegacyConversionUnavailable):
                processor.format_document(str(source), str(Path(tmpdir) / "legacy_formatted.docx"))


class LegacyFormatTests(unittest.TestCase):
    def test_format_document_handles_doc_via_soffice_when_available(self):
        converter = SofficeConverter()
        if not converter.available:
            self.skipTest("LibreOffice soffice is not available")

        with tempfile.TemporaryDirectory(prefix="wfp_legacy_doc_test_") as tmpdir:
            root = Path(tmpdir)
            docx_source = root / "legacy_source.docx"
            source_doc = Document()
            source_doc.add_paragraph("旧格式标题")
            source_doc.add_paragraph("第一段正文")
            source_doc.save(docx_source)

            legacy_dir = root / "legacy"
            legacy_dir.mkdir()
            profile = root / "lo_profile"
            profile.mkdir()
            proc = subprocess.run(
                [
                    converter.soffice_path,
                    "--headless",
                    f"-env:UserInstallation={profile.as_uri()}",
                    "--convert-to",
                    "doc",
                    "--outdir",
                    str(legacy_dir),
                    str(docx_source),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=120,
            )
            generated = sorted(legacy_dir.glob("*.doc"))
            if proc.returncode != 0 or not generated:
                detail = (proc.stderr or proc.stdout or "LibreOffice did not generate a .doc file").strip()
                self.skipTest(f"LibreOffice .doc export unavailable: {detail}")

            output = root / "legacy_formatted.docx"
            processor = WordProcessor(DEFAULT_CONFIG.copy(), soffice_path=converter.soffice_path)
            try:
                processor.format_document(str(generated[0]), str(output))
                self.assertTrue(output.exists())
                formatted = Document(output)
                self.assertTrue(any(paragraph.text.strip() for paragraph in formatted.paragraphs))
            finally:
                processor._cleanup_temp_files()


class EndToEndFormatTests(unittest.TestCase):
    def test_format_document_handles_direct_formats(self):
        with tempfile.TemporaryDirectory(prefix="wfp_format_test_") as tmpdir:
            root = Path(tmpdir)
            sources = []

            txt_source = root / "sample.txt"
            txt_source.write_text("测试标题\n\n第一段正文\n1. 小标题\n第二段正文", encoding="utf-8")
            sources.append(txt_source)

            md_source = root / "sample.md"
            md_source.write_text("# Markdown标题\n\n**第一段**\n\n1. 步骤一", encoding="utf-8")
            sources.append(md_source)

            docx_source = root / "sample.docx"
            source_doc = Document()
            source_doc.add_paragraph("Word标题")
            source_doc.add_paragraph("第一段正文")
            source_doc.save(docx_source)
            sources.append(docx_source)

            for source in sources:
                processor = WordProcessor(DEFAULT_CONFIG.copy())
                output = root / f"{source.stem}_formatted.docx"
                try:
                    processor.format_document(str(source), str(output))
                    self.assertTrue(output.exists(), f"missing output for {source.suffix}")
                    formatted = Document(output)
                    texts = [paragraph.text for paragraph in formatted.paragraphs]
                    self.assertTrue(any(text.strip() for text in texts), f"empty output for {source.suffix}")
                finally:
                    processor._cleanup_temp_files()


def main(argv=None):
    suite = unittest.defaultTestLoader.loadTestsFromModule(__import__(__name__))
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(suite)
    if result.wasSuccessful():
        print("所有单元测试通过")
        return 0
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
