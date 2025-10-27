import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, Menu
import json
import os
import re
import logging
import shutil
import win32com.client
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from tkinterdnd2 import DND_FILES, TkinterDnD

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class WordProcessor:
    def __init__(self, config, log_callback=None):
        self.config = config
        self.temp_files = []
        self.log_callback = log_callback
        # --- MODIFICATION: Attribute to hold the COM application instance ---
        self.com_app = None

    def _log(self, message):
        if self.log_callback: self.log_callback(message)

    def _cleanup_temp_files(self):
        self._log("正在清理本轮临时文件...")
        for f in self.temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
                    self._log(f"  > 临时文件 {os.path.basename(f)} 已删除")
            except OSError as e:
                self._log(f"  > 警告：删除临时文件 {f} 失败: {e}")
        self.temp_files.clear()

    # --- MODIFICATION: Lazily get or create a single COM app instance ---
    def _get_wps_app(self):
        if self.com_app is None:
            self._log("首次需要，正在启动WPS/Word应用...")
            try:
                self.com_app = win32com.client.Dispatch('KWPS.Application')
                self._log("  > 已成功连接到WPS。")
            except Exception:
                try:
                    self.com_app = win32com.client.Dispatch('Word.Application')
                    self._log("  > 已成功连接到Word。")
                except Exception as e:
                    raise RuntimeError(f"未能启动WPS或Word，请确保已安装。错误: {e}")
            self.com_app.Visible = False
        return self.com_app
        
    # --- MODIFICATION: New method to quit the app at the very end ---
    def quit_com_app(self):
        if self.com_app:
            self._log("所有任务完成，正在关闭WPS/Word应用...")
            self.com_app.Quit()
            self.com_app = None
            self._log("  > 应用已关闭。")

    def convert_to_docx(self, input_path):
        file_ext = os.path.splitext(input_path)[1].lower()
        is_from_txt = (file_ext == '.txt')
        temp_dir = os.path.dirname(input_path)
        base_name = os.path.splitext(os.path.basename(input_path))[0]

        if file_ext == '.docx':
            self._log("检测到 .docx 文件，正在创建安全的处理副本...")
            temp_docx_path = os.path.join(temp_dir, f"~temp_copy_{base_name}.docx")
            shutil.copy2(input_path, temp_docx_path)
            self.temp_files.append(temp_docx_path)
            self._log(f"  > 副本创建成功: {os.path.basename(temp_docx_path)}")
            return temp_docx_path, is_from_txt

        temp_docx_path = os.path.join(temp_dir, f"~temp_converted_{base_name}.docx")
        self.temp_files.append(temp_docx_path)

        if file_ext == '.txt':
            self._log("检测到 .txt 文件，正在创建 .docx...")
            doc = Document()
            try:
                with open(input_path, 'r', encoding='utf-8') as f:
                    for line in f: doc.add_paragraph(line.strip())
                self._log("  > 已使用 UTF-8 编码读取TXT文件。")
            except UnicodeDecodeError:
                self._log("  > UTF-8读取失败，尝试使用 GBK 编码...")
                with open(input_path, 'r', encoding='gbk') as f:
                    for line in f: doc.add_paragraph(line.strip())
                self._log("  > 已成功使用 GBK 编码读取TXT文件。")
            doc.save(temp_docx_path)
            self._log("TXT转换完成。")
            return temp_docx_path, is_from_txt
        elif file_ext in ['.wps', '.doc']:
            self._log(f"正在转换 {file_ext} 文件为 .docx...")
            app = self._get_wps_app()
            doc_com = app.Documents.Open(os.path.abspath(input_path), ReadOnly=1)
            doc_com.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=12)
            doc_com.Close()
            self._log("文件格式转换完成。")
            return temp_docx_path, is_from_txt
        
        raise ValueError(f"不支持的文件格式: {file_ext}")

    def _preprocess_com_tasks(self, docx_path):
        self._log("正在对副本执行预处理（接受所有修订、转换自动编号）...")
        app = self._get_wps_app()
        try:
            doc_com = app.Documents.Open(os.path.abspath(docx_path))
            
            doc_com.TrackRevisions = False; self._log("  > 已关闭修订追踪。")
            
            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions(); self._log("  > 已接受文档副本中的所有修订。")
            
            doc_com.Content.ListFormat.ConvertNumbersToText(); self._log("  > 已将副本中的自动编号转换为文本。")
            
            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions(); self._log("  > 已接受编号转换产生的修订。")
            
            doc_com.TrackRevisions = False
            
            doc_com.Save()
            doc_com.Close()
            self._log("预处理完成。")
        except Exception as e:
            self._log(f"警告：执行预处理任务时出错: {e}")

    def _create_page_number(self, paragraph, text):
        font_name = self.config['page_number_font']
        font_size = self.config['page_number_size']
        self._set_run_font(paragraph.add_run('— '), font_name, font_size, set_color=True)
        run_field = paragraph.add_run()
        self._set_run_font(run_field, font_name, font_size, set_color=True)
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = text
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run_field._r.extend([fldChar1, instrText, fldChar2])
        self._set_run_font(paragraph.add_run(' —'), font_name, font_size, set_color=True)

    def _apply_page_setup(self, doc):
        self._log("正在应用页面边距和页码设置...")
        for section in doc.sections:
            section.top_margin = Cm(self.config['margin_top'])
            section.bottom_margin = Cm(self.config['margin_bottom'])
            section.left_margin = Cm(self.config['margin_left'])
            section.right_margin = Cm(self.config['margin_right'])
            section.footer_distance = Cm(self.config['footer_distance'])

            if self.config['page_number_align'] == '居中':
                p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
                p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; self._create_page_number(p, 'PAGE')
            elif self.config['page_number_align'] == '奇偶分页':
                doc.settings.odd_and_even_pages_header_footer = True
                footer_odd = section.footer
                p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
                p_odd.clear(); p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT; self._create_page_number(p_odd, 'PAGE')
                
                footer_even = section.even_page_footer
                p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
                p_even.clear(); p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT; self._create_page_number(p_even, 'PAGE')

    def _set_run_font(self, run, font_name, size_pt, set_color=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        if set_color: run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def _apply_font_to_runs(self, para, font_name, size_pt, set_color=False):
        for run in para.runs: self._set_run_font(run, font_name, size_pt, set_color=set_color)

    def _strip_leading_whitespace(self, para):
        if not para.runs: return
        while para.runs and not para.runs[0].text.strip():
            p = para._p
            p.remove(para.runs[0]._r)
        if not para.runs: return
        first_run = para.runs[0]
        original_text = first_run.text
        stripped_text = original_text.lstrip()
        if original_text != stripped_text:
            first_run.text = stripped_text
            self._log("  > 已移除段落前的多余空格。")
    
    def _reset_pagination_properties(self, para):
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_lines_together = False
        para.paragraph_format.page_break_before = False
        para.paragraph_format.keep_together = False 

    def _format_heading(self, para, level, is_from_txt): 
        if self.config['set_outline'] and not is_from_txt: 
            try: para.style = f'Heading {level}'
            except KeyError:
                try: para.style = f'标题 {level}'
                except KeyError: self._log(f"  > 警告: 样式 'Heading {level}' 或 '标题 {level}' 未找到。")

    def _apply_text_indent_and_align(self, para):
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.left_indent = Cm(self.config['left_indent_cm'])
        para.paragraph_format.right_indent = Cm(self.config['right_indent_cm'])
        ind = para._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn("w:firstLineChars"), "200")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _iter_block_items(self, parent):
        parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P): yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl): yield Table(child, parent)
    
    def _find_title_paragraph_index(self, doc, is_from_txt):
        ch_num = r'[一二三四五六七八九十百千万零]+'
        re_h1 = re.compile(r'^' + ch_num + r'\s*、')
        re_h2 = re.compile(r'^[（\(]' + ch_num + r'[）\)]')

        if is_from_txt:
            self._log("文档源自 TXT，采用智能规则查找题目...")
            for idx, block in enumerate(self._iter_block_items(doc)):
                if isinstance(block, Paragraph) and block.text.strip():
                    text_to_check = block.text.strip()
                    if re_h1.match(text_to_check) or re_h2.match(text_to_check):
                        self._log(f"  > 首个非空行 (块 {idx + 1}) 符合标题格式，认定本文档无独立题目。")
                        return -1
                    else:
                        self._log(f"  > 在块 {idx + 1} 发现首个非空段落，认定为题目。")
                        return idx
            self._log("  > 扫描结束，未找到任何非空段落。"); return -1
        else:
            self._log("正在预扫描以确定居中题目位置...")
            for idx, block in enumerate(self._iter_block_items(doc)):
                if not isinstance(block, Paragraph) or not block.text.strip(): continue
                para = block; text_to_check = para.text.lstrip()
                if re_h1.match(text_to_check) or re_h2.match(text_to_check):
                    self._log("  > 发现一级/二级标题，在此之前未找到居中题目。"); return -1
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    self._log(f"  > 在块 {idx + 1} 发现潜在题目。"); return idx
            self._log("  > 扫描结束，未能在主要标题前找到任何居中段落作为题目。"); return -1

    def format_document(self, input_path, output_path):
        processing_path, is_from_txt = self.convert_to_docx(input_path)
        if not is_from_txt: self._preprocess_com_tasks(processing_path)
        
        doc = Document(processing_path)
        
        if not is_from_txt:
            self._log("正在重置标题样式以保留原文粗体/斜体...")
            for i in range(1, 5):
                for style_name_tpl in [f'Heading {i}', f'标题 {i}']:
                    try:
                        style = doc.styles[style_name_tpl]
                        style.font.bold = None; style.font.italic = None
                        self._log(f"  > 样式 '{style_name_tpl}' 的强制粗体/斜体已重置。")
                        break 
                    except KeyError: continue

        all_blocks = list(self._iter_block_items(doc))
        processed_indices = set()
        apply_color = not is_from_txt

        if not is_from_txt:
            self._log("正在扫描图表标题...")
            for idx, block in enumerate(all_blocks):
                is_pic_para = isinstance(block, Paragraph) and ('<w:drawing>' in block._p.xml or '<w:pict>' in block._p.xml)
                is_table = isinstance(block, Table)
                if not (is_pic_para or is_table): continue
                for direction in [-1, 1]:
                    caption_found = False
                    for i in range(idx + direction, -1 if direction == -1 else len(all_blocks), direction):
                        if i in processed_indices: continue
                        potential_caption = all_blocks[i]
                        if not isinstance(potential_caption, Paragraph): break 
                        text = potential_caption.text.strip()
                        if text: 
                            if potential_caption.alignment == WD_ALIGN_PARAGRAPH.CENTER and (text.startswith("图") or text.startswith("表")):
                                detected_type = "图" if text.startswith("图") else "表"
                                self._log(f"  > 发现 {detected_type} 的标题: \"{text[:30]}...\" (在段落 {i+1})")
                                config_font_key = f'{("figure" if detected_type == "图" else "table")}_caption_font'
                                config_size_key = f'{("figure" if detected_type == "图" else "table")}_caption_size'
                                self._apply_font_to_runs(potential_caption, self.config[config_font_key], self.config[config_size_key], set_color=apply_color)
                                processed_indices.add(i)
                                caption_found = True
                            break 
                    if caption_found: break 

        title_block_index = self._find_title_paragraph_index(doc, is_from_txt)
        if title_block_index != -1: processed_indices.add(title_block_index)

        self._log("预扫描完成，开始逐段格式化...")
        re_h1 = re.compile(r'^[一二三四五六七八九十百千万零]+\s*、')
        re_h2 = re.compile(r'^[（\(][一二三四五六七八九十百千万零]+[）\)]')
        re_h3 = re.compile(r'^\d+\s*[\.．]')
        re_h4 = re.compile(r'^[（\(]\d+[）\)]')
        re_attachment = re.compile(r'^附件\s*(\d+|[一二三四五六七八九十百千万零]+)?\s*[:：]?$')

        if title_block_index != -1:
            para = all_blocks[title_block_index]
            self._log(f"段落 {title_block_index + 1}: 题目 - \"{para.text[:30]}...\"")
            self._strip_leading_whitespace(para)
            self._apply_font_to_runs(para, self.config['title_font'], self.config['title_size'], set_color=apply_color)
            para.alignment, para.paragraph_format.first_line_indent = WD_ALIGN_PARAGRAPH.CENTER, None
            self._reset_pagination_properties(para)

        block_idx = 0
        while block_idx < len(all_blocks):
            block = all_blocks[block_idx]
            
            if block_idx in processed_indices:
                if block_idx != title_block_index: self._log(f"块 {block_idx + 1}: 已作为图表/附件标题处理 - 跳过")
                block_idx += 1; continue

            current_block_num = block_idx + 1
            if isinstance(block, Table): 
                self._log(f"块 {current_block_num}: 表格 - 跳过"); block_idx += 1; continue
            
            para = block
            if not para.text.strip(): 
                self._log(f"段落 {current_block_num}: 空白 - 跳过"); block_idx += 1; continue
            
            is_pic = '<w:drawing>' in para._p.xml or '<w:pict>' in para._p.xml
            is_embedded_obj = '<w:object>' in para._p.xml
            if is_pic or is_embedded_obj:
                self._log(f"段落 {current_block_num}: {'图片' if is_pic else '附件'} - 仅格式化文字")
                text_to_check = para.text.lstrip(); para_text_preview = text_to_check[:30].replace("\n", " ")
                if re_h1.match(text_to_check):
                    self._log(f"  > 文字识别为一级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color)
                elif re_h2.match(text_to_check):
                    self._log(f"  > 文字识别为二级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                elif text_to_check:
                    self._log(f"  > 文字识别为正文: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                block_idx += 1; continue

            original_text, text_to_check_stripped = para.text, para.text.strip()
            text_to_check = para.text.lstrip()
            leading_space_count = len(original_text) - len(text_to_check)
            para_text_preview = text_to_check[:30].replace("\n", " ")
            
            spacing = para._p.get_or_add_pPr().get_or_add_spacing()
            spacing.set(qn('w:beforeAutospacing'), '0'); spacing.set(qn('w:afterAutospacing'), '0')
            para.paragraph_format.space_before, para.paragraph_format.space_after = Pt(0), Pt(0)
            para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])

            is_attachment_enabled = self.config.get('enable_attachment_formatting', False)
            is_attachment_candidate = False
            if is_from_txt:
                if re_attachment.match(text_to_check_stripped): is_attachment_candidate = True
            elif para.alignment in [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY, None] and re_attachment.match(text_to_check_stripped):
                is_attachment_candidate = True

            if is_attachment_enabled and is_attachment_candidate:
                self._log(f"段落 {current_block_num}: 附件标识 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['attachment_font'], self.config['attachment_size'], set_color=apply_color)
                self._reset_pagination_properties(para)
                para.paragraph_format.page_break_before = True
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.left_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                attachment_title_idx, search_idx = -1, block_idx + 1
                while search_idx < len(all_blocks):
                    next_block = all_blocks[search_idx]
                    if isinstance(next_block, Paragraph) and next_block.text.strip():
                        attachment_title_idx = search_idx; break
                    elif isinstance(next_block, Table): break
                    search_idx += 1

                if attachment_title_idx != -1:
                    para_title = all_blocks[attachment_title_idx]
                    self._log(f"  > 识别到附件标题: \"{para_title.text.strip()[:30]}...\" (在段落 {attachment_title_idx + 1})")
                    processed_indices.add(attachment_title_idx)
                    self._strip_leading_whitespace(para_title)
                    self._apply_font_to_runs(para_title, self.config['title_font'], self.config['title_size'], set_color=apply_color)
                    para_title.alignment, para_title.paragraph_format.first_line_indent = WD_ALIGN_PARAGRAPH.CENTER, None
                    self._reset_pagination_properties(para_title)
                
                block_idx = (attachment_title_idx + 1) if attachment_title_idx != -1 else search_idx
                continue
            
            elif re_h1.match(text_to_check):
                self._log(f"段落 {current_block_num}: 一级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para); self._format_heading(para, 1, is_from_txt)
                self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color); self._apply_text_indent_and_align(para); self._reset_pagination_properties(para)
            
            elif re_h2.match(text_to_check):
                self._log(f"段落 {current_block_num}: 二级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                parts = para.text.split('。', 1)
                if len(parts) == 2 and parts[1].strip():
                    self._log("  > 检测到二级标题与正文在同一段落，执行段内格式拆分。")
                    title_len = len(parts[0]) + 1
                    original_runs_info = [{'text': r.text, 'bold': r.bold, 'italic': r.italic, 'underline': r.underline, 'font_color': r.font.color.rgb} for r in para.runs]
                    para.clear()
                    char_count = 0
                    for run_info in original_runs_info:
                        run_text = run_info['text']; run_end_pos = char_count + len(run_text)
                        
                        title_run, body_run = None, None
                        if run_end_pos <= title_len:
                            new_run = para.add_run(run_text)
                            self._set_run_font(new_run, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                        elif char_count >= title_len:
                            new_run = para.add_run(run_text)
                            self._set_run_font(new_run, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                        else:
                            split_index = title_len - char_count
                            title_part, body_part = run_text[:split_index], run_text[split_index:]
                            if title_part:
                                title_run = para.add_run(title_part)
                                self._set_run_font(title_run, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                            if body_part:
                                body_run = para.add_run(body_part)
                                self._set_run_font(body_run, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                        
                        runs_to_format = [r for r in [title_run, body_run] if r] or [para.runs[-1]]
                        for r in runs_to_format:
                            r.bold = run_info['bold']; r.italic = run_info['italic']; r.underline = run_info['underline']
                            if run_info['font_color']: r.font.color.rgb = run_info['font_color']
                        
                        char_count = run_end_pos
                    self._format_heading(para, 2, is_from_txt); self._apply_text_indent_and_align(para); self._reset_pagination_properties(para)
                else:
                    match = re.match(r'^[（\(](.+?)[）\)](.*)', text_to_check, re.DOTALL)
                    if match and not (text_to_check.startswith('（') and text_to_check.strip().endswith('）')):
                        self._log("  > 已将二级标题的括号统一为中文括号。")
                        for r in para.runs: r.text = r.text.replace('(', '（', 1).replace(')', '）', 1)
                    self._format_heading(para, 2, is_from_txt); self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color); self._apply_text_indent_and_align(para); self._reset_pagination_properties(para)
            
            elif re_h3.match(text_to_check) or re_h4.match(text_to_check):
                level = 3 if re_h3.match(text_to_check) else 4
                self._log(f"段落 {current_block_num}: {'三' if level == 3 else '四'}级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para); self._format_heading(para, level, is_from_txt)
                self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color); self._apply_text_indent_and_align(para); self._reset_pagination_properties(para)
            elif not is_from_txt and (para.alignment in [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT] or leading_space_count > 5 or (para.paragraph_format.first_line_indent is None or para.paragraph_format.first_line_indent.pt == 0) and leading_space_count == 0):
                if para.alignment in [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]: self._log(f"段落 {current_block_num}: {'居中' if para.alignment == WD_ALIGN_PARAGRAPH.CENTER else '右对齐'}正文 - 保留原对齐")
                elif leading_space_count > 5: self._log(f"段落 {current_block_num}: 正文 (保留前导空格) - \"{para_text_preview}...\"")
                else: self._log(f"段落 {current_block_num}: 正文 (保留0缩进) - \"{para_text_preview}...\"")
                self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                para.alignment = para.alignment if para.alignment is not None else WD_ALIGN_PARAGRAPH.LEFT
                self._reset_pagination_properties(para)
            else:
                log_reason = "源自TXT，强制缩进" if is_from_txt else "应用标准缩进"
                self._log(f"段落 {current_block_num}: 正文 ({log_reason}) - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                self._apply_text_indent_and_align(para); self._reset_pagination_properties(para)
            
            block_idx += 1
        
        self._apply_page_setup(doc)
        self._log("正在保存最终文档...")
        doc.save(output_path)


class WordFormatterGUI:
    def __init__(self, master):
        self.master = master
        master.title("Word文档智能排版工具 v2.6.0")
        master.geometry("1200x800")

        self.font_size_map = {
            '一号 (26pt)': 26, '小一 (24pt)': 24, '二号 (22pt)': 22, '小二 (18pt)': 18,
            '三号 (16pt)': 16, '小三 (15pt)': 15, '四号 (14pt)': 14, '小四 (12pt)': 12,
            '五号 (10.5pt)': 10.5, '小五 (9pt)': 9
        }
        self.font_size_map_rev = {v: k for k, v in self.font_size_map.items()}
        
        self.default_params = {
            'margin_top': 3.7, 'margin_bottom': 3.5, 'margin_left': 2.8, 'margin_right': 2.6,
            'footer_distance': 2.5, 'line_spacing': 28, 'page_number_align': '奇偶分页',
            'title_font': '方正小标宋简体', 'h1_font': '黑体', 'h2_font': '楷体_GB2312', 'body_font': '仿宋_GB2312',
            'page_number_font': '宋体', 'table_caption_font': '黑体', 'figure_caption_font': '黑体', 'attachment_font': '黑体',
            'title_size': 22, 'h1_size': 16, 'h2_size': 16, 'body_size': 16, 'page_number_size': 14,
            'table_caption_size': 14, 'figure_caption_size': 14, 'attachment_size': 16,
            'left_indent_cm': 0.0, 'right_indent_cm': 0.0,
            'set_outline': True, 'enable_attachment_formatting': True
        }
        self.font_options = {
            'title': ['方正小标宋简体', '方正小标宋_GBK', '华文中宋', '宋体'], 'h1': ['黑体', '方正黑体_GBK', '方正黑体简体', '华文黑体', '宋体'],
            'h2': ['楷体_GB2312', '方正楷体_GBK', '楷体', '方正楷体简体', '华文楷体', '宋体'],
            'body': ['仿宋_GB2312', '方正仿宋_GBK', '仿宋', '方正仿宋简体', '华文仿宋', '宋体'], 'page_number': ['宋体', 'Times New Roman'],
            'table_caption': ['黑体', '宋体', '仿宋_GB2312'], 'figure_caption': ['黑体', '宋体', '仿宋_GB2312'], 'attachment': ['黑体', '宋体', '仿宋_GB2312']
        }
        self.set_outline_var = tk.BooleanVar(value=self.default_params['set_outline'])
        self.enable_attachment_var = tk.BooleanVar(value=self.default_params['enable_attachment_formatting'])
        self.entries = {}
        
        self.default_config_path = "default_config.json"
        
        self.create_menu()
        self.create_widgets()
        self.load_initial_config()

    def create_menu(self):
        menubar = Menu(self.master)
        help_menu = Menu(menubar, tearoff=0)
        help_menu.add_command(label="使用说明", command=self.show_help_window)
        menubar.add_cascade(label="帮助", menu=help_menu)
        self.master.config(menu=menubar)

    def create_widgets(self):
        main_pane = ttk.PanedWindow(self.master, orient=tk.HORIZONTAL)
        main_pane.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left_frame = ttk.Frame(main_pane, padding=5); main_pane.add(left_frame, weight=1)
        notebook = ttk.Notebook(left_frame); notebook.pack(fill=tk.BOTH, expand=True); self.notebook = notebook

        file_tab = ttk.Frame(notebook); notebook.add(file_tab, text=' 文件批量处理 ')
        list_frame = ttk.LabelFrame(file_tab, text="待处理文件列表（可拖拽文件或文件夹）")
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
        self.file_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, selectmode=tk.EXTENDED)
        scrollbar.config(command=self.file_listbox.yview); scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.file_listbox.drop_target_register(DND_FILES); self.file_listbox.dnd_bind('<<Drop>>', self.handle_drop)
        self.placeholder_label = ttk.Label(self.file_listbox, text="可以拖拽文件或文件夹到这里", foreground="grey")
        
        file_button_frame = ttk.Frame(file_tab); file_button_frame.pack(fill=tk.X, pady=5)
        ttk.Button(file_button_frame, text="添加文件", command=self.add_files).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(file_button_frame, text="添加文件夹", command=self.add_folder).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(file_button_frame, text="移除文件", command=self.remove_files).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(file_button_frame, text="清空列表", command=self.clear_list).pack(side=tk.LEFT, expand=True, fill=tk.X)

        text_tab = ttk.Frame(notebook); notebook.add(text_tab, text=' 直接输入文本 ')
        text_frame = ttk.LabelFrame(text_tab, text="在此处输入或粘贴文本")
        text_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.direct_text_input = scrolledtext.ScrolledText(text_frame, height=10, wrap=tk.WORD)
        self.direct_text_input.pack(fill=tk.BOTH, expand=True)

        right_frame = ttk.Frame(main_pane, padding=5); main_pane.add(right_frame, weight=2)
        params_frame = ttk.LabelFrame(right_frame, text="参数设置"); params_frame.pack(fill=tk.X, pady=(0, 5))
        params_frame.columnconfigure(1, weight=1); params_frame.columnconfigure(3, weight=1); params_frame.columnconfigure(5, weight=1)
        
        def create_entry(label, var, r, c): ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=2); e = ttk.Entry(params_frame); e.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=2); self.entries[var] = e
        def create_combo(label, var, opts, r, c): ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=2); cb = ttk.Combobox(params_frame, values=opts, state='readonly'); cb.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=2); self.entries[var] = cb
        def create_font_combo(label, var, opts, r, c): ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=2); cb = ttk.Combobox(params_frame, values=opts); cb.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=2); self.entries[var] = cb
        def create_font_size_combo(label, var, r, c): ttk.Label(params_frame, text=label).grid(row=r, column=c, sticky=tk.W, padx=5, pady=2); cb = ttk.Combobox(params_frame, values=list(self.font_size_map.keys())); cb.grid(row=r, column=c+1, sticky=tk.EW, padx=5, pady=2); self.entries[var] = cb
        
        row = 0
        create_combo("页码对齐", 'page_number_align', ['奇偶分页', '居中'], row, 0); create_font_combo("题目字体", 'title_font', self.font_options['title'], row, 2); create_font_size_combo("题目字号", 'title_size', row, 4); row+=1
        create_entry("页脚距(cm)", 'footer_distance', row, 0); create_font_combo("一级标题字体", 'h1_font', self.font_options['h1'], row, 2); create_font_size_combo("一级标题字号", 'h1_size', row, 4); row+=1
        create_entry("行间距(磅)", 'line_spacing', row, 0); create_font_combo("二级标题字体", 'h2_font', self.font_options['h2'], row, 2); create_font_size_combo("二级标题字号", 'h2_size', row, 4); row+=1
        create_entry("段落左缩进(cm)", 'left_indent_cm', row, 0); create_font_combo("正文/三四级字体", 'body_font', self.font_options['body'], row, 2); create_font_size_combo("正文/三四级字号", 'body_size', row, 4); row+=1
        create_entry("段落右缩进(cm)", 'right_indent_cm', row, 0); create_font_combo("页码字体", 'page_number_font', self.font_options['page_number'], row, 2); create_font_size_combo("页码字号", 'page_number_size', row, 4); row+=1
        create_entry("上边距(cm)", 'margin_top', row, 0); create_font_combo("表格标题字体", 'table_caption_font', self.font_options['table_caption'], row, 2); create_font_size_combo("表格标题字号", 'table_caption_size', row, 4); row+=1
        create_entry("下边距(cm)", 'margin_bottom', row, 0); create_font_combo("图形标题字体", 'figure_caption_font', self.font_options['figure_caption'], row, 2); create_font_size_combo("图形标题字号", 'figure_caption_size', row, 4); row+=1
        create_entry("左边距(cm)", 'margin_left', row, 0); create_entry("右边距(cm)", 'margin_right', row, 2); row+=1

        ttk.Separator(params_frame, orient='horizontal').grid(row=row, column=0, columnspan=6, sticky='ew', pady=5); row+=1
        ttk.Checkbutton(params_frame, text="附件设置 (段前分页、识别标题)", variable=self.enable_attachment_var).grid(row=row, column=0, columnspan=2, sticky=tk.W, padx=5, pady=2)
        create_font_combo("附件标识字体", 'attachment_font', self.font_options['attachment'], row, 2); create_font_size_combo("附件标识字号", 'attachment_size', row, 4); row+=1
        
        ttk.Checkbutton(params_frame, text="自动设置大纲级别 (对非TXT源文件)", variable=self.set_outline_var).grid(row=row, columnspan=6, pady=5); row+=1
        
        log_frame = ttk.LabelFrame(right_frame, text="调试日志"); log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.debug_text = scrolledtext.ScrolledText(log_frame, height=8, state='disabled', wrap=tk.WORD); self.debug_text.pack(fill=tk.BOTH, expand=True)
        
        button_frame = ttk.Frame(right_frame); button_frame.pack(fill=tk.X, pady=5)
        ttk.Button(button_frame, text="加载配置", command=self.load_config).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(button_frame, text="保存配置", command=self.save_config).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(button_frame, text="保存为默认", command=self.save_default_config).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(button_frame, text="恢复内置默认", command=self.load_defaults).pack(side=tk.LEFT, expand=True, fill=tk.X)

        style = ttk.Style(); style.configure('Success.TButton', font=('Helvetica', 10, 'bold'), foreground='green')
        ttk.Button(right_frame, text="开始排版", style='Success.TButton', command=self.start_processing).pack(fill=tk.X, ipady=8, pady=5)

        self._update_listbox_placeholder()

    def log_to_debug_window(self, message):
        self.master.update_idletasks(); self.debug_text.config(state='normal'); self.debug_text.insert(tk.END, message + '\n'); self.debug_text.config(state='disabled'); self.debug_text.see(tk.END)
    
    def load_initial_config(self):
        if os.path.exists(self.default_config_path):
            try:
                with open(self.default_config_path, 'r', encoding='utf-8') as f: config = json.load(f)
                self._apply_config(config); self.log_to_debug_window(f"已加载默认配置文件: {self.default_config_path}")
            except Exception as e:
                self.log_to_debug_window(f"加载默认配置 '{self.default_config_path}' 失败: {e}。将使用内置默认值。"); self.load_defaults()
        else: self.log_to_debug_window("未找到默认配置文件，将使用内置默认值。"); self.load_defaults()

    def _apply_config(self, config):
        self.set_outline_var.set(config.get('set_outline', True))
        self.enable_attachment_var.set(config.get('enable_attachment_formatting', True))
        for key, value in config.items():
            if key in ['set_outline', 'enable_attachment_formatting']: continue
            widget = self.entries.get(key)
            if widget:
                if "_size" in key: widget.set(self.font_size_map_rev.get(value, str(value)))
                elif isinstance(widget, ttk.Combobox): widget.set(value)
                else: widget.delete(0, tk.END); widget.insert(0, str(value))

    def load_defaults(self):
        self.set_outline_var.set(self.default_params['set_outline'])
        self.enable_attachment_var.set(self.default_params['enable_attachment_formatting'])
        for key, value in self.default_params.items():
            if key in ['set_outline', 'enable_attachment_formatting']: continue
            widget = self.entries.get(key)
            if "_size" in key: widget.set(self.font_size_map_rev.get(value, str(value)))
            elif isinstance(widget, ttk.Combobox): widget.set(value)
            else: widget.delete(0, tk.END); widget.insert(0, str(value))
    
    def collect_config(self):
        config = {}
        for key, widget in self.entries.items():
            value = widget.get().strip()
            if "_size" in key:
                config[key] = self.font_size_map.get(value, 16)
            else:
                try: config[key] = float(value) if '.' in value else int(value)
                except (ValueError, TypeError): config[key] = value
        config['set_outline'] = self.set_outline_var.get()
        config['enable_attachment_formatting'] = self.enable_attachment_var.get()
        return config

    def save_config(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON files", "*.json")])
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f: json.dump(self.collect_config(), f, ensure_ascii=False, indent=4)
            messagebox.showinfo("成功", f"配置已保存至 {file_path}")
    
    def save_default_config(self):
        try:
            with open(self.default_config_path, 'w', encoding='utf-8') as f: json.dump(self.collect_config(), f, ensure_ascii=False, indent=4)
            messagebox.showinfo("成功", f"当前配置已保存为默认配置。\n下次启动软件时将自动加载。")
        except Exception as e: messagebox.showerror("错误", f"保存默认配置失败: {e}")

    def load_config(self):
        file_path = filedialog.askopenfilename(filetypes=[("JSON files", "*.json")])
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f: loaded_config = json.load(f)
                self._apply_config(loaded_config); messagebox.showinfo("成功", "配置已加载")
            except Exception as e: messagebox.showerror("错误", f"加载配置文件失败: {e}")

    def _update_listbox_placeholder(self):
        if self.file_listbox.size() == 0: self.placeholder_label.place(in_=self.file_listbox, relx=0.5, rely=0.5, anchor=tk.CENTER)
        else: self.placeholder_label.place_forget()

    def handle_drop(self, event): self._add_paths_to_listbox(self.master.tk.splitlist(event.data))

    def _add_paths_to_listbox(self, paths):
        current_files = set(self.file_listbox.get(0, tk.END)); added_count = 0
        for path in paths:
            if os.path.isdir(path):
                for root, _, files in os.walk(path):
                    for f in files:
                        if f.lower().endswith(('.docx', '.doc', '.wps', '.txt')):
                            full_path = os.path.join(root, f)
                            if full_path not in current_files: self.file_listbox.insert(tk.END, full_path); current_files.add(full_path); added_count += 1
            elif os.path.isfile(path) and path.lower().endswith(('.docx', '.doc', '.wps', '.txt')):
                if path not in current_files: self.file_listbox.insert(tk.END, path); current_files.add(path); added_count += 1
        if added_count > 0: self.log_to_debug_window(f"通过按钮或拖拽添加了 {added_count} 个新文件。")
        self._update_listbox_placeholder()

    def add_files(self):
        files = filedialog.askopenfilenames(filetypes=[("所有支持的文件", "*.docx;*.doc;*.wps;*.txt"), ("Word 文档", "*.docx;*.doc"), ("WPS 文档", "*.wps"), ("纯文本", "*.txt")])
        if files: self._add_paths_to_listbox(files)
        
    def add_folder(self):
        folder = filedialog.askdirectory()
        if folder: self._add_paths_to_listbox([folder])

    def remove_files(self):
        selected_indices = self.file_listbox.curselection()
        if not selected_indices: messagebox.showinfo("提示", "请先在列表中选择要移除的文件。"); return
        for index in sorted(selected_indices, reverse=True): self.file_listbox.delete(index)
        self._update_listbox_placeholder()

    def clear_list(self): self.file_listbox.delete(0, tk.END); self._update_listbox_placeholder()

    def show_help_window(self):
        help_win = tk.Toplevel(self.master); help_win.title("使用说明"); help_win.geometry("600x550")
        help_text_widget = scrolledtext.ScrolledText(help_win, wrap=tk.WORD, state='disabled')
        help_text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        help_content = """
Word文档智能排版工具 v2.6.0 - 使用说明

本工具旨在提供一键式的专业文档排版体验，支持批量处理和高度自定义。

【核心功能模式】
1. 文件批量处理：可拖拽或添加 .docx, .doc, .wps, .txt 文件。
2. 直接输入文本：直接粘贴文本进行排版。

【操作流程】
1. 选择模式并添加内容。
2. （可选）在“参数设置”区调整格式。
3. 点击“开始排版”，并选择输出位置。

【参数与配置】
- 自定义字体：所有字体设置项都支持手动输入。
- 配置管理：可通过按钮“恢复内置默认”、“保存/加载配置”、“保存为默认”来管理排版方案。

【智能识别特性】
- 自动识别题目、1-4级标题、图/表标题。
- 附件处理：
  - 自动识别“附件1”、“附件一”等标识行。
  - 可在参数区独立设置附件标识的字体、字号。
  - 启用“附件设置”后，会自动为附件添加“段前分页”，并将附件标识后的第一段文字识别为附件的独立标题。
- 保留原文格式：统一格式时，会保留【加粗、斜体】等。
- 二级标题智能拆分：若二级标题后紧跟正文（如“（一）标题。正文...”），会自动在【同一个段落内】为标题和正文应用不同格式。
- 豁免内容：表格、图片、附件等内容会自动跳过格式化。

【安全提示】
本工具【绝对不会】修改您的任何原始文件。所有操作都在后台的临时副本上进行，确保源文件100%安全。
"""
        help_text_widget.config(state='normal')
        help_text_widget.insert('1.0', help_content.strip().replace('   -', '\t-'))
        help_text_widget.config(state='disabled')

    def start_processing(self):
        if not messagebox.askokcancel("处理前重要提示", "为防止数据丢失，请在继续前关闭所有已打开的Word和WPS程序。\n\n本程序在转换文件格式时需调用Word/WPS，可能会导致您未保存的工作被强制关闭。\n\n您确定要继续吗？"):
            self.log_to_debug_window("用户已取消操作。"); return
            
        self.debug_text.config(state='normal'); self.debug_text.delete('1.0', tk.END); self.debug_text.config(state='disabled')
        
        processor = WordProcessor(self.collect_config(), self.log_to_debug_window)
        active_tab_index = self.notebook.index(self.notebook.select())

        try:
            if active_tab_index == 0:
                file_list = self.file_listbox.get(0, tk.END)
                if not file_list: messagebox.showwarning("警告", "文件列表为空！"); return
                output_dir = filedialog.askdirectory(title="请选择一个文件夹用于存放处理后的文件")
                if not output_dir: return

                success_count, fail_count = 0, 0
                for i, input_path in enumerate(file_list):
                    try:
                        self.log_to_debug_window(f"\n--- 开始处理文件 {i+1}/{len(file_list)}: {os.path.basename(input_path)} ---")
                        base_name = os.path.splitext(os.path.basename(input_path))[0]
                        output_path = os.path.join(output_dir, f"{base_name}_formatted.docx")
                        processor.format_document(input_path, output_path)
                        self.log_to_debug_window(f"✅ 文件处理成功，已保存至: {output_path}")
                        success_count += 1
                    except Exception as e:
                        logging.error(f"处理文件失败: {input_path}\n{e}", exc_info=True)
                        self.log_to_debug_window(f"\n❌ 处理文件 {os.path.basename(input_path)} 时发生严重错误：\n{e}"); fail_count += 1
                    finally: processor._cleanup_temp_files()
                
                summary_message = f"批量处理完成！\n\n成功: {success_count}个\n失败: {fail_count}个"
                if fail_count > 0: summary_message += "\n\n失败详情请查看日志窗口。"
                messagebox.showinfo("完成", summary_message); self.log_to_debug_window(f"\n🎉 {summary_message}")

            elif active_tab_index == 1:
                text_content = self.direct_text_input.get('1.0', tk.END).strip()
                if not text_content: messagebox.showwarning("警告", "文本框内容为空！"); return
                output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile="formatted_document.docx")
                if not output_path: return
                
                temp_file_path = None
                try:
                    fd, temp_file_path = tempfile.mkstemp(suffix=".txt", text=True)
                    with os.fdopen(fd, 'w', encoding='utf-8') as tmp: tmp.write(text_content)
                    
                    self.log_to_debug_window(f"\n--- 开始处理输入的文本 ---")
                    processor.format_document(temp_file_path, output_path)
                    self.log_to_debug_window("\n🎉 排版全部完成！")
                    messagebox.showinfo("完成", f"文档排版成功！\n文件已保存至：\n{output_path}")
                finally:
                    processor._cleanup_temp_files()
                    if temp_file_path and os.path.exists(temp_file_path):
                        try: os.remove(temp_file_path); self.log_to_debug_window(f"  > 输入文本的临时文件已删除")
                        except OSError: pass
        except Exception as e:
            logging.error(f"处理过程中发生严重错误: {e}", exc_info=True)
            self.log_to_debug_window(f"\n❌ 处理过程中发生严重错误：\n{e}")
            messagebox.showerror("错误", f"处理过程中发生错误：\n{e}")
        finally:
            processor.quit_com_app()

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = WordFormatterGUI(root)
    root.mainloop()