#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word Formatter Pro CLI and agent entry point.

Formats doc/docx/wps/txt/md files into standardized docx files. The formatting
rules are provided by wfp_core.py, extracted from Word-Formatter-Pro v2.6.9.
"""

from __future__ import annotations

import argparse
import atexit
import glob
import json
import logging
import os
from pathlib import Path
import platform
import shutil
import subprocess
import sys
import tempfile
import threading
from dataclasses import dataclass

try:
    import win32com.client
except ImportError:
    win32com = None

from docx import Document

from wfp_core import WordProcessor as WfpWordProcessor


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

logger = logging.getLogger("wfp_cli")

CONFIG_FILE_NAME = "wfp_config.json"
SUPPORTED_EXTENSIONS = {".doc", ".docx", ".wps", ".txt", ".md"}
LEGACY_WORD_EXTENSIONS = {".doc", ".wps"}

DEFAULT_CONFIG = {
    "page_number_align": "奇偶分页",
    "footer_distance": 2.5,
    "line_spacing": 28,
    "margin_top": 3.7,
    "margin_bottom": 3.5,
    "margin_left": 2.8,
    "margin_right": 2.6,
    "title_font": "方正小标宋简体",
    "h1_font": "黑体",
    "h2_font": "楷体_GB2312",
    "body_font": "仿宋_GB2312",
    "page_number_font": "宋体",
    "table_caption_font": "黑体",
    "figure_caption_font": "黑体",
    "attachment_font": "黑体",
    "subtitle_font": "楷体_GB2312",
    "title_size": 22,
    "h1_size": 16,
    "h2_size": 16,
    "body_size": 16,
    "page_number_size": 14,
    "table_caption_size": 14,
    "figure_caption_size": 14,
    "attachment_size": 16,
    "subtitle_size": 16,
    "title_line_spacing": 33,
    "subtitle_line_spacing": 33,
    "left_indent_cm": 0.0,
    "right_indent_cm": 0.0,
    "set_outline": True,
    "enable_attachment_formatting": True,
    "force_a4": False,
    "use_custom_english_font": False,
    "english_font": "Times New Roman",
    "remove_blank_lines": True,
    "normalize_punctuation": False,
    "enable_table_formatting": False,
    "table_header_font": "仿宋_GB2312",
    "table_font": "仿宋_GB2312",
    "table_size": 12,
    "table_line_spacing": 22,
    "table_row_height_cm": 0.7,
    "table_auto_col_width": True,
    "table_width_percent": 100,
    "table_header_bold": True,
    "table_smart_align": False,
    "table_unified_borders": True,
    "table_border_size_pt": 0.5,
    "table_col_min_pct": 8,
    "table_col_max_pct": 45,
    "table_short_text_len": 4,
}

FONT_SIZE_NAMES = {
    26: "一号",
    24: "小一",
    22: "二号",
    18: "小二",
    16: "三号",
    15: "小三",
    14: "四号",
    12: "小四",
    10.5: "五号",
    9: "小五",
}

CONFIG_DESCRIPTIONS = {
    "page_number_align": ("页码对齐", "奇偶分页 / 居中"),
    "footer_distance": ("页脚距离", "厘米"),
    "line_spacing": ("正文行距", "磅"),
    "margin_top": ("上边距", "厘米"),
    "margin_bottom": ("下边距", "厘米"),
    "margin_left": ("左边距", "厘米"),
    "margin_right": ("右边距", "厘米"),
    "title_font": ("题目字体", "例如 方正小标宋简体"),
    "h1_font": ("一级标题字体", "一、二、三、"),
    "h2_font": ("二级标题字体", "（一）（二）"),
    "body_font": ("正文字体", "正文、三四级标题默认字体"),
    "page_number_font": ("页码字体", "页脚页码字体"),
    "table_caption_font": ("表格标题字体", "表 1 等标题"),
    "figure_caption_font": ("图形标题字体", "图 1 等标题"),
    "attachment_font": ("附件标识字体", "附件：等段落"),
    "subtitle_font": ("副标题字体", "题目下方副标题"),
    "title_size": ("题目字号", "pt"),
    "h1_size": ("一级标题字号", "pt"),
    "h2_size": ("二级标题字号", "pt"),
    "body_size": ("正文字号", "pt"),
    "page_number_size": ("页码字号", "pt"),
    "table_caption_size": ("表格标题字号", "pt"),
    "figure_caption_size": ("图形标题字号", "pt"),
    "attachment_size": ("附件标识字号", "pt"),
    "subtitle_size": ("副标题字号", "pt"),
    "title_line_spacing": ("题目行距", "磅"),
    "subtitle_line_spacing": ("副标题行距", "磅"),
    "left_indent_cm": ("段落左缩进", "厘米"),
    "right_indent_cm": ("段落右缩进", "厘米"),
    "set_outline": ("设置大纲级别", "true/false"),
    "enable_attachment_formatting": ("附件格式化", "true/false"),
    "force_a4": ("强制 A4", "true/false"),
    "use_custom_english_font": ("单独设置数字和字母字体", "true/false"),
    "english_font": ("数字和字母字体", "默认 Times New Roman"),
    "remove_blank_lines": ("删除 TXT/MD 多余空行", "true/false"),
    "normalize_punctuation": ("符号标准化", "true/false，保守修复中英文标点混用"),
    "enable_table_formatting": ("表格内容自动调整", "true/false"),
    "table_header_font": ("表头字体", "enable_table_formatting 为 true 时生效"),
    "table_font": ("表格正文字体", "enable_table_formatting 为 true 时生效"),
    "table_size": ("表格字号", "pt"),
    "table_line_spacing": ("表格行距", "磅"),
    "table_row_height_cm": ("表格行高", "厘米"),
    "table_auto_col_width": ("自动调整列宽", "true/false"),
    "table_width_percent": ("表格宽度", "百分比"),
    "table_header_bold": ("表头加粗", "true/false"),
    "table_smart_align": ("表格智能对齐", "true/false"),
    "table_unified_borders": ("统一表格边框", "true/false"),
    "table_border_size_pt": ("表格边框粗细", "pt"),
    "table_col_min_pct": ("自动列宽最小比例", "百分比"),
    "table_col_max_pct": ("自动列宽最大比例", "百分比"),
    "table_short_text_len": ("表格短文本判定长度", "字符数"),
}

NUMBERING_WARNING = (
    "提示：本次至少有一个 Word/WPS 类文档未能确认完成自动编号转文本。"
    "自动编号的文本、段落可能没有正常格式化，请打开输出文件自行检查编号字体、字号和段落缩进。"
)


class UserActionRequired(RuntimeError):
    """Raised when the user must manually convert a file before formatting."""


class BackendUnavailableError(RuntimeError):
    """Raised when a document conversion backend cannot be started."""


class WindowsOfficeBackend:
    """Use WPS or Microsoft Word COM automation on Windows."""

    def __init__(self, prog_ids: tuple[str, ...]):
        self.prog_ids = prog_ids
        self.app = None
        self.app_prog_id = None

    @property
    def available(self) -> bool:
        return platform.system() == "Windows" and win32com is not None

    def _get_app(self, prog_id: str):
        if not self.available:
            raise BackendUnavailableError("Windows WPS/Office COM 接口不可用。")
        if self.app is not None and self.app_prog_id == prog_id:
            return self.app
        if self.app is not None:
            self.close()

        logger.info("正在启动 %s ...", prog_id)
        try:
            self.app = win32com.client.Dispatch(prog_id)
            self.app_prog_id = prog_id
            self.app.Visible = False
            try:
                self.app.DisplayAlerts = 0
            except Exception:
                pass
            logger.info("已连接到 %s。", prog_id)
            return self.app
        except Exception as exc:
            self.app = None
            self.app_prog_id = None
            raise BackendUnavailableError(f"未能启动 {prog_id}: {exc}") from exc

    def close(self):
        if self.app is None:
            return
        try:
            logger.info("正在关闭 %s ...", self.app_prog_id or "WPS/Word")
            self.app.Quit()
        except Exception as exc:
            logger.warning("关闭 %s 时出错，已忽略: %s", self.app_prog_id or "WPS/Word", exc)
        finally:
            self.app = None
            self.app_prog_id = None

    def convert_to_docx(self, input_path: Path, output_path: Path) -> Path:
        errors = []
        for prog_id in self.prog_ids:
            doc_com = None
            try:
                app = self._get_app(prog_id)
                doc_com = app.Documents.Open(str(input_path.resolve()), ReadOnly=1)
                doc_com.SaveAs2(str(output_path.resolve()), FileFormat=12)
                logger.info("文件格式转换完成（%s）。", prog_id)
                return output_path
            except Exception as exc:
                errors.append(f"{prog_id}: {exc}")
                logger.warning("%s 转换失败: %s", prog_id, exc)
            finally:
                if doc_com is not None:
                    try:
                        doc_com.Close(False)
                    except Exception:
                        pass

        raise UserActionRequired(
            f"无法将 {input_path} 转为 docx。请先用 WPS 或 Microsoft Word 将文档另存为 .docx 后再输入。"
            f"详细错误: {' | '.join(errors)}"
        )

    def preprocess_numbering(self, docx_path: Path) -> bool:
        errors = []
        for prog_id in self.prog_ids:
            doc_com = None
            try:
                app = self._get_app(prog_id)
                doc_com = app.Documents.Open(str(docx_path.resolve()))
                doc_com.TrackRevisions = False
                if doc_com.Revisions.Count > 0:
                    doc_com.AcceptAllRevisions()
                doc_com.Content.ListFormat.ConvertNumbersToText()
                if doc_com.Revisions.Count > 0:
                    doc_com.AcceptAllRevisions()
                doc_com.TrackRevisions = False
                doc_com.Save()
                logger.info("已将自动编号转换为文本（%s）。", prog_id)
                return True
            except Exception as exc:
                errors.append(f"{prog_id}: {exc}")
                logger.warning("%s 自动编号预处理失败，准备尝试下一个后端: %s", prog_id, exc)
            finally:
                if doc_com is not None:
                    try:
                        doc_com.Close(False)
                    except Exception:
                        pass
        logger.warning("已跳过自动编号转换: %s", " | ".join(errors))
        return False


class LibreOfficeBackend:
    """Use LibreOffice headless mode for doc/wps conversion."""

    def __init__(self, soffice_path: str | None = None, timeout: int = 120):
        self.timeout = timeout
        self.soffice_path = soffice_path or self._find_soffice()
        self._lock = threading.Lock()
        self._profile_dir = None
        if self.soffice_path:
            self._profile_dir = tempfile.mkdtemp(prefix="wfp_libre_profile_")
            atexit.register(shutil.rmtree, self._profile_dir, ignore_errors=True)

    @property
    def available(self) -> bool:
        return self.soffice_path is not None

    @staticmethod
    def _find_soffice() -> str | None:
        executable = "soffice.com" if platform.system() == "Windows" else "soffice"
        found = shutil.which(executable)
        if found:
            return found

        common_paths = []
        if platform.system() == "Darwin":
            common_paths.extend(
                [
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                    "/opt/homebrew/bin/soffice",
                    "/usr/local/bin/soffice",
                ]
            )
        elif platform.system() == "Windows":
            for base in (os.environ.get("PROGRAMFILES"), os.environ.get("PROGRAMFILES(X86)")):
                if base:
                    common_paths.extend(
                        [
                            os.path.join(base, "LibreOffice", "program", "soffice.com"),
                            os.path.join(base, "LibreOffice", "program", "soffice.exe"),
                        ]
                    )
        else:
            common_paths.extend(
                [
                    "/usr/bin/soffice",
                    "/usr/local/bin/soffice",
                    "/snap/bin/libreoffice",
                    "/opt/libreoffice/program/soffice",
                ]
            )

        for path in common_paths:
            if path and os.path.isfile(path):
                return path
        return None

    def convert_to_docx(self, input_path: Path, output_path: Path) -> Path:
        if not self.available:
            raise UserActionRequired(
                f"LibreOffice (soffice) 不可用，无法转换 {input_path}。"
                "请安装 LibreOffice，或先将文档转为 .docx 后再输入。"
            )

        with self._lock:
            temp_dir = Path(tempfile.mkdtemp(prefix="wfp_libre_out_"))
            try:
                profile = str(Path(self._profile_dir).resolve()).replace(os.sep, "/")
                cmd = [
                    self.soffice_path,
                    "--headless",
                    "--norestore",
                    f"-env:UserInstallation=file:///{profile}",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(temp_dir),
                    str(input_path.resolve()),
                ]
                creationflags = 0
                if platform.system() == "Windows" and hasattr(subprocess, "CREATE_NO_WINDOW"):
                    creationflags = subprocess.CREATE_NO_WINDOW
                proc = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    creationflags=creationflags,
                )
                try:
                    stdout, stderr = proc.communicate(timeout=self.timeout)
                except subprocess.TimeoutExpired as exc:
                    proc.kill()
                    proc.communicate()
                    raise RuntimeError(f"LibreOffice 转换超时 ({self.timeout}s): {input_path}") from exc

                if proc.returncode != 0:
                    detail = stderr.decode("utf-8", errors="ignore") or stdout.decode("utf-8", errors="ignore")
                    raise UserActionRequired(
                        f"LibreOffice 无法将 {input_path} 转为 docx。请先手动转为 .docx 后再输入。详细错误: {detail}"
                    )

                generated_files = list(temp_dir.glob("*.docx"))
                if not generated_files:
                    raise UserActionRequired(
                        f"LibreOffice 未生成 {input_path} 对应的 docx。请先手动转为 .docx 后再输入。"
                    )
                output_path.parent.mkdir(parents=True, exist_ok=True)
                shutil.move(str(generated_files[0]), str(output_path))
                logger.info("LibreOffice 文件格式转换完成。")
                return output_path
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)

    def preprocess_numbering(self, docx_path: Path) -> bool:
        logger.info("LibreOffice 后端不支持自动编号转文本，已跳过: %s", docx_path)
        return False

    def close(self):
        if self._profile_dir:
            shutil.rmtree(self._profile_dir, ignore_errors=True)
            self._profile_dir = None


class CliWordProcessor(WfpWordProcessor):
    """Headless adapter around the v2.6.9 formatting engine."""

    def __init__(
        self,
        config: dict,
        backend=None,
        log_callback=None,
        remove_blank_lines: bool = True,
        convert_numbering: bool = True,
    ):
        super().__init__(config, log_callback=log_callback, remove_blank_lines=remove_blank_lines)
        self.backend = backend
        self.convert_numbering = convert_numbering
        self.numbering_attention_needed = False
        self._work_dir = Path(tempfile.mkdtemp(prefix="wfp_cli_"))

    def _log(self, message):
        if self.log_callback:
            self.log_callback(message)
        else:
            logger.debug(message)

    def close(self):
        self._cleanup_temp_files()
        shutil.rmtree(self._work_dir, ignore_errors=True)

    def _new_temp_docx_path(self, input_path: Path, prefix: str) -> Path:
        safe_stem = input_path.stem.replace(" ", "_") or "document"
        fd, raw_path = tempfile.mkstemp(prefix=f"{prefix}_{safe_stem}_", suffix=".docx", dir=self._work_dir)
        os.close(fd)
        path = Path(raw_path)
        self.temp_files.append(str(path))
        return path

    def convert_to_docx(self, input_path):
        input_path = Path(input_path)
        file_ext = input_path.suffix.lower()
        is_from_txt = file_ext in (".txt", ".md")

        if file_ext == ".docx":
            self._log("检测到 .docx 文件，正在创建安全的处理副本...")
            temp_docx_path = self._new_temp_docx_path(input_path, "copy")
            shutil.copy2(input_path, temp_docx_path)
            return str(temp_docx_path), False

        temp_docx_path = self._new_temp_docx_path(input_path, "converted")

        if file_ext == ".txt":
            self._log("检测到 .txt 文件，正在创建 .docx...")
            text_content = self._read_text_file(str(input_path))
            if self.remove_blank_lines:
                text_content = self._remove_blank_lines_from_text(text_content)
                self._log("  > 已删除 TXT 中的多余空行。")
            doc = Document()
            for line in text_content.split("\n"):
                doc.add_paragraph(line.strip())
            doc.save(temp_docx_path)
            return str(temp_docx_path), is_from_txt

        if file_ext == ".md":
            self._log("检测到 .md 文件，正在清理 Markdown 标记并创建 .docx...")
            raw_text = self._read_text_file(str(input_path))
            cleaned_text = self._clean_markdown(raw_text)
            if self.remove_blank_lines:
                cleaned_text = self._remove_blank_lines_from_text(cleaned_text)
                self._log("  > 已删除 Markdown 文本中的多余空行。")
            doc = Document()
            for line in cleaned_text.split("\n"):
                doc.add_paragraph(line.strip())
            doc.save(temp_docx_path)
            return str(temp_docx_path), is_from_txt

        if file_ext in LEGACY_WORD_EXTENSIONS:
            self._log(f"正在转换 {file_ext} 文件为 .docx...")
            if self.backend is None:
                raise UserActionRequired(
                    f"处理 {input_path} 需要 WPS/Office 或 LibreOffice。"
                    "请安装可用后端，或先将文档转为 .docx 后再输入。"
                )
            try:
                temp_docx_path.unlink()
            except FileNotFoundError:
                pass
            self.backend.convert_to_docx(input_path, temp_docx_path)
            return str(temp_docx_path), is_from_txt

        raise ValueError(f"不支持的文件格式: {file_ext}")

    def _preprocess_com_tasks(self, docx_path):
        if not self.convert_numbering:
            self.numbering_attention_needed = True
            self._log("已按当前平台或参数设置跳过自动编号转换。")
            return
        if self.backend is None:
            self.numbering_attention_needed = True
            self._log("未配置可用后端，已跳过自动编号转换。")
            return
        self._log("正在对副本执行预处理（接受所有修订、转换自动编号）...")
        converted = self.backend.preprocess_numbering(Path(docx_path))
        if not converted:
            self.numbering_attention_needed = True
            self._log("自动编号转换未执行或未成功；最终文件需人工检查编号字体字号。")


@dataclass
class InputRecord:
    source: Path
    relative: Path


@dataclass
class Job:
    source: Path
    output: Path


def parse_value(raw: str):
    text = raw.strip()
    lowered = text.lower()
    if lowered in {"true", "yes", "on"}:
        return True
    if lowered in {"false", "no", "off"}:
        return False
    if lowered in {"null", "none"}:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return raw


def apply_set_overrides(config: dict, set_items: list[str] | None):
    if not set_items:
        return
    for item in set_items:
        if "=" not in item:
            raise ValueError(f"--set 参数必须使用 key=value 形式: {item}")
        key, raw_value = item.split("=", 1)
        key = key.strip()
        if key not in DEFAULT_CONFIG:
            raise KeyError(f"未知配置项: {key}")
        config[key] = parse_value(raw_value)


def apply_convenience_overrides(config: dict, args):
    if getattr(args, "enable_table_formatting", False):
        config["enable_table_formatting"] = True
    if getattr(args, "disable_table_formatting", False):
        config["enable_table_formatting"] = False
    if getattr(args, "enable_custom_english_font", False):
        config["use_custom_english_font"] = True
    if getattr(args, "disable_custom_english_font", False):
        config["use_custom_english_font"] = False
    if getattr(args, "english_font", None):
        config["english_font"] = args.english_font
        config["use_custom_english_font"] = True
    if getattr(args, "normalize_punctuation", False):
        config["normalize_punctuation"] = True
    if getattr(args, "disable_normalize_punctuation", False):
        config["normalize_punctuation"] = False


def normalize_config(config: dict) -> dict:
    if "use_custom_english_font" not in config and config.get("use_times_new_roman"):
        config["use_custom_english_font"] = True
        config.setdefault("english_font", "Times New Roman")
    merged = dict(DEFAULT_CONFIG)
    merged.update(config)
    for key, default_value in DEFAULT_CONFIG.items():
        value = merged.get(key)
        if value is None or (isinstance(value, str) and not value.strip()):
            merged[key] = default_value
    return merged


def load_json_file(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError(f"配置文件必须是 JSON 对象: {path}")
    return data


def load_config(config_path=None, config_json=None):
    source = "内置默认配置"
    config = dict(DEFAULT_CONFIG)

    if config_path:
        path = Path(config_path).expanduser().resolve()
        config.update(load_json_file(path))
        source = str(path)
    else:
        cwd_config = Path.cwd() / CONFIG_FILE_NAME
        if cwd_config.exists():
            config.update(load_json_file(cwd_config))
            source = str(cwd_config.resolve())

    if config_json:
        inline_config = json.loads(config_json)
        if not isinstance(inline_config, dict):
            raise ValueError("--config-json 必须是 JSON 对象")
        config.update(inline_config)
        source += " + --config-json"

    return normalize_config(config), source


def show_config(args):
    config, source = load_config(args.config, args.config_json)
    apply_set_overrides(config, args.set)
    apply_convenience_overrides(config, args)
    payload = {
        "config_source": source,
        "config_file_auto_load": str((Path.cwd() / CONFIG_FILE_NAME).resolve()),
        "supported_inputs": sorted(SUPPORTED_EXTENSIONS),
        "output": "单文件输出 .docx；多文件或目录输出到目录，目录输入会递归保留原目录结构。",
        "font_size_names": {str(key): value for key, value in FONT_SIZE_NAMES.items()},
        "optional_features": [
            "enable_table_formatting=true 启用表格内容自动调整",
            "use_custom_english_font=true 并设置 english_font，可单独指定数字和字母字体，默认 Times New Roman",
            "normalize_punctuation=true 启用符号标准化",
        ],
        "config": {},
    }
    for key, value in config.items():
        name, note = CONFIG_DESCRIPTIONS.get(key, (key, ""))
        payload["config"][key] = {"value": value, "name": name, "note": note}
    print(json.dumps(payload, ensure_ascii=False, indent=2))


def save_config(args):
    config, source = load_config(args.config, args.config_json)
    apply_set_overrides(config, args.set)
    apply_convenience_overrides(config, args)
    output_path = Path(args.output or CONFIG_FILE_NAME).expanduser().resolve()
    output_path.write_text(json.dumps(config, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(str(output_path))
    print(
        "提示：配置已保存。还可按需启用表格内容自动调整、自定义数字和字母字体"
        "（默认 Times New Roman）或符号标准化；修改后可再次运行 format 重新转换。",
        file=sys.stderr,
    )
    logger.info("配置来源: %s", source)


def libreoffice_install_help():
    print(
        """LibreOffice 安装命令参考：
macOS (Homebrew):
  brew install --cask libreoffice

Debian/Ubuntu:
  sudo apt-get update
  sudo apt-get install libreoffice

Fedora:
  sudo dnf install libreoffice

Arch Linux:
  sudo pacman -S libreoffice-fresh

安装后请确认 soffice 可用：
  soffice --version
"""
    )


def is_supported_file(path: Path) -> bool:
    return path.is_file() and not path.name.startswith("~") and path.suffix.lower() in SUPPORTED_EXTENSIONS


def collect_records(paths: list[str], recursive: bool = True) -> list[InputRecord]:
    if not paths:
        raise ValueError("请提供至少一个输入文件或目录。")

    resolved_inputs = [Path(path).expanduser().resolve() for path in paths]
    include_root_prefix = len(resolved_inputs) > 1
    records = []

    for input_path in resolved_inputs:
        if input_path.is_file():
            if not is_supported_file(input_path):
                raise ValueError(f"不支持的文件格式: {input_path}")
            records.append(InputRecord(input_path, Path(input_path.name)))
            continue

        if input_path.is_dir():
            iterator = input_path.rglob("*") if recursive else input_path.glob("*")
            for file_path in sorted(iterator):
                if not is_supported_file(file_path):
                    continue
                rel = file_path.relative_to(input_path)
                if include_root_prefix:
                    rel = Path(input_path.name) / rel
                records.append(InputRecord(file_path, rel))
            continue

        raise FileNotFoundError(f"输入路径不存在: {input_path}")

    if not records:
        raise FileNotFoundError("未找到可处理的文件。支持格式: doc, docx, wps, txt, md")
    return records


def formatted_relative_path(relative: Path) -> Path:
    return relative.with_name(f"{relative.stem}_formatted.docx")


def unique_path(path: Path, seen: set[Path]) -> Path:
    candidate = path
    counter = 2
    while candidate in seen:
        candidate = path.with_name(f"{path.stem}_{counter}{path.suffix}")
        counter += 1
    seen.add(candidate)
    return candidate


def build_jobs(input_paths: list[str], output_arg: str | None, recursive: bool = True) -> list[Job]:
    records = collect_records(input_paths, recursive=recursive)
    output_path = Path(output_arg).expanduser().resolve() if output_arg else None

    if len(records) == 1 and records[0].source.is_file():
        source = records[0].source
        if output_path is None:
            output = source.with_name(f"{source.stem}_formatted.docx")
        elif output_path.suffix.lower() == ".docx":
            output = output_path
        else:
            output = output_path / f"{source.stem}_formatted.docx"
        return [Job(source, output)]

    if output_path is None:
        first = Path(input_paths[0]).expanduser().resolve()
        if len(input_paths) == 1 and first.is_dir():
            output_dir = first.parent / f"{first.name}_formatted"
        else:
            output_dir = Path.cwd() / "wfp_formatted"
    else:
        if output_path.suffix.lower() == ".docx":
            raise ValueError("多文件或目录输入时，--output 必须是输出目录，不能是单个 .docx 文件。")
        output_dir = output_path

    jobs = []
    seen = set()
    for record in records:
        rel_output = unique_path(formatted_relative_path(record.relative), seen)
        jobs.append(Job(record.source, output_dir / rel_output))
    return jobs


def office_prog_ids(office_app: str) -> tuple[str, ...]:
    if office_app == "wps":
        return ("KWPS.Application",)
    if office_app == "word":
        return ("Word.Application",)
    return ("KWPS.Application", "Word.Application")


def build_backend(args):
    backend_name = args.backend
    if backend_name == "none":
        return None
    if backend_name == "office":
        backend = WindowsOfficeBackend(office_prog_ids(args.office_app))
        if not backend.available:
            raise BackendUnavailableError("Office/WPS 后端不可用：需要 Windows、pywin32、WPS 或 Microsoft Word。")
        return backend
    if backend_name == "libreoffice":
        return LibreOfficeBackend(args.soffice, args.timeout)

    if platform.system() == "Windows":
        backend = WindowsOfficeBackend(office_prog_ids(args.office_app))
        if backend.available:
            return backend
        logger.warning("Windows Office/WPS COM 接口不可用；doc/wps 转换和自动编号转换将不可用。")
        return None
    return LibreOfficeBackend(args.soffice, args.timeout)


def add_config_override_args(parser):
    parser.add_argument("--config", help="JSON 配置文件路径；未指定时自动读取当前目录 wfp_config.json")
    parser.add_argument("--config-json", help="内联 JSON 配置对象，会覆盖配置文件中的同名字段")
    parser.add_argument("--set", action="append", help="覆盖单个配置项，格式 key=value，可重复使用")
    parser.add_argument("--enable-table-formatting", action="store_true", help="启用表格内容自动调整")
    parser.add_argument("--disable-table-formatting", action="store_true", help="关闭表格内容自动调整")
    parser.add_argument("--enable-custom-english-font", action="store_true", help="启用数字和字母单独字体")
    parser.add_argument("--disable-custom-english-font", action="store_true", help="关闭数字和字母单独字体")
    parser.add_argument("--english-font", help="数字和字母字体；设置后会自动启用 use_custom_english_font")
    parser.add_argument("--normalize-punctuation", action="store_true", help="启用符号标准化")
    parser.add_argument("--disable-normalize-punctuation", action="store_true", help="关闭符号标准化")


def format_command(args):
    input_paths = []
    if args.inputs:
        input_paths.extend(args.inputs)
    input_paths.extend(args.paths or [])

    config, config_source = load_config(args.config, args.config_json)
    apply_set_overrides(config, args.set)
    apply_convenience_overrides(config, args)
    logger.info("使用配置: %s", config_source)

    remove_blank_lines = False if args.keep_blank_lines else bool(config.get("remove_blank_lines", True))
    convert_numbering = args.convert_numbering
    if convert_numbering is None:
        convert_numbering = platform.system() == "Windows"

    jobs = build_jobs(input_paths, args.output, recursive=not args.no_recursive)
    backend = build_backend(args)
    processor = CliWordProcessor(
        config,
        backend=backend,
        remove_blank_lines=remove_blank_lines,
        convert_numbering=convert_numbering,
    )

    success_count = 0
    fail_count = 0
    output_paths = []
    try:
        for index, job in enumerate(jobs, start=1):
            try:
                logger.info("开始处理 %s/%s: %s", index, len(jobs), job.source)
                job.output.parent.mkdir(parents=True, exist_ok=True)
                processor.format_document(str(job.source), str(job.output))
                processor._cleanup_temp_files()
                output_paths.append(job.output.resolve())
                success_count += 1
            except Exception as exc:
                fail_count += 1
                processor._cleanup_temp_files()
                logger.error("处理失败: %s: %s", job.source, exc, exc_info=args.verbose)
    finally:
        processor.close()
        if backend and hasattr(backend, "close"):
            backend.close()

    for output_path in output_paths:
        print(str(output_path))

    if processor.numbering_attention_needed:
        print(NUMBERING_WARNING, file=sys.stderr)
    logger.info("处理完成：成功 %s 个，失败 %s 个。", success_count, fail_count)
    return 1 if fail_count else 0


def build_parser():
    parser = argparse.ArgumentParser(
        description="Word Formatter Pro CLI：将 doc/docx/wps/txt/md 按公文排版规则格式化为 docx。"
    )
    subparsers = parser.add_subparsers(dest="command", required=True)

    fmt = subparsers.add_parser("format", help="格式化单文件、多文件或目录")
    fmt.add_argument("paths", nargs="*", help="输入文件或目录，可一次传入多个")
    fmt.add_argument("-i", "--input", dest="inputs", action="append", help="输入文件或目录，可重复")
    fmt.add_argument("-o", "--output", help="输出文件或目录；目录输入会保留原目录结构")
    add_config_override_args(fmt)
    fmt.add_argument("--backend", choices=["auto", "office", "libreoffice", "none"], default="auto")
    fmt.add_argument("--office-app", choices=["auto", "wps", "word"], default="auto")
    fmt.add_argument("--soffice", help="LibreOffice soffice 可执行文件路径")
    fmt.add_argument("--timeout", type=int, default=120, help="LibreOffice 单文件转换超时秒数")
    fmt.add_argument("--no-recursive", action="store_true", help="目录输入时不递归子目录；默认递归")
    numbering = fmt.add_mutually_exclusive_group()
    numbering.add_argument("--convert-numbering", dest="convert_numbering", action="store_true", default=None)
    numbering.add_argument("--skip-numbering", dest="convert_numbering", action="store_false")
    fmt.add_argument("--keep-blank-lines", action="store_true", help="保留 TXT/MD 原始空行")
    fmt.add_argument("-v", "--verbose", action="store_true", help="显示详细日志")

    show = subparsers.add_parser("show-config", help="显示当前配置、默认配置说明和可选增强项")
    add_config_override_args(show)

    save = subparsers.add_parser("save-config", help="保存配置到当前目录 wfp_config.json 或指定路径")
    add_config_override_args(save)
    save.add_argument("-o", "--output", help="输出 JSON 配置文件路径，默认 ./wfp_config.json")

    subparsers.add_parser("install-help", help="显示 LibreOffice 安装提示")
    return parser


def configure_logging(verbose: bool = False):
    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        stream=sys.stderr,
    )


def main(argv=None):
    parser = build_parser()
    args = parser.parse_args(argv)
    configure_logging(getattr(args, "verbose", False))

    if args.command == "show-config":
        show_config(args)
        return 0
    if args.command == "save-config":
        save_config(args)
        return 0
    if args.command == "install-help":
        libreoffice_install_help()
        return 0
    if args.command == "format":
        return format_command(args)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except UserActionRequired as exc:
        configure_logging(False)
        logger.error(str(exc))
        raise SystemExit(1)
    except Exception as exc:
        configure_logging(False)
        logger.error(str(exc))
        if "LibreOffice" in str(exc):
            libreoffice_install_help()
        raise SystemExit(1)
