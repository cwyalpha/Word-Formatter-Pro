#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Headless Word Formatter Pro core extracted from 2.6.9/wfp.py.

This file intentionally contains only the document formatting engine. The GUI
entry point and Tkinter dependencies from the original script are not included.
"""

import os
import re
import logging
import shutil
import tempfile

try:
    import win32com.client
except ImportError:  # Non-Windows agents can still use the python-docx rules.
    win32com = None

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

class WordProcessor:
    def __init__(self, config, log_callback=None, remove_blank_lines=True):
        self.config = config
        self.temp_files = []
        self.log_callback = log_callback
        self.com_app = None
        self.remove_blank_lines = remove_blank_lines

    def _log(self, message):
        if self.log_callback: self.log_callback(message)

    def _cleanup_temp_files(self):
        self._log("正在清理本轮临时文件...")
        for f in self.temp_files:
            try:
                if os.path.exists(f):
                    os.remove(f)
                    self._log(f"  > 临时文件 {os.path.basename(f)} 已删除")
            except OSError as e:
                self._log(f"  > 警告：删除临时文件 {f} 失败: {e}")
        self.temp_files.clear()

    def _get_wps_app(self):
        if self.com_app is None:
            self._log("首次需要，正在启动WPS/Word应用...")
            try:
                self.com_app = win32com.client.Dispatch('KWPS.Application')
                self._log("  > 已成功连接到WPS。")
            except Exception:
                try:
                    self.com_app = win32com.client.Dispatch('Word.Application')
                    self._log("  > 已成功连接到Word。")
                except Exception as e:
                    raise RuntimeError(f"未能启动WPS或Word，请确保已安装。错误: {e}")
            self.com_app.Visible = False
        return self.com_app
        
    def quit_com_app(self):
        if self.com_app:
            self._log("所有任务完成，正在关闭WPS/Word应用...")
            self.com_app.Quit()
            self.com_app = None
            self._log("  > 应用已关闭。")

    # ------------------------------------------------------------------
    # Conservative Chinese punctuation normalization
    # ------------------------------------------------------------------
    @staticmethod
    def _has_chinese(text):
        return bool(re.search(r'[\u4e00-\u9fff]', text or ''))

    @staticmethod
    def _is_digit_or_latin(ch):
        if not ch:
            return False
        code = ord(ch)
        return (
            ch.isdigit()
            or ('A' <= ch <= 'Z')
            or ('a' <= ch <= 'z')
            or (0xFF10 <= code <= 0xFF19)
            or (0xFF21 <= code <= 0xFF3A)
            or (0xFF41 <= code <= 0xFF5A)
        )

    @staticmethod
    def _prev_non_space(text, index):
        i = index - 1
        while i >= 0 and text[i].isspace():
            i -= 1
        return text[i] if i >= 0 else ''

    @staticmethod
    def _next_non_space(text, index):
        i = index + 1
        while i < len(text) and text[i].isspace():
            i += 1
        return text[i] if i < len(text) else ''

    @classmethod
    def _is_after_digit_or_latin(cls, text, index):
        return cls._is_digit_or_latin(cls._prev_non_space(text, index))

    @classmethod
    def _is_before_digit_or_latin(cls, text, index):
        return cls._is_digit_or_latin(cls._next_non_space(text, index))

    @classmethod
    def _normalize_ellipsis(cls, text):
        dot_chars = {'.', '\uff0e', '\u3002'}
        chars = []
        i = 0
        while i < len(text):
            ch = text[i]
            if ch not in dot_chars:
                chars.append(ch)
                i += 1
                continue

            j = i + 1
            while j < len(text) and text[j] in dot_chars:
                j += 1

            if j - i >= 3 and not cls._is_after_digit_or_latin(text, i):
                chars.append('\u2026\u2026')
            else:
                chars.append(text[i:j])
            i = j

        return ''.join(chars)

    @classmethod
    def _normalize_simple_punctuation(cls, text):
        replacements = {
            ',': '\uff0c',
            '.': '\u3002',
            '\uff0e': '\u3002',
            ';': '\uff1b',
            ':': '\uff1a',
            '?': '\uff1f',
            '!': '\uff01',
        }
        chars = list(text)
        for i, ch in enumerate(chars):
            if ch not in replacements:
                continue
            if ch in {'.', '\uff0e', '\u3002'}:
                prev_is_dot = i > 0 and text[i - 1] in {'.', '\uff0e', '\u3002'}
                next_is_dot = i + 1 < len(text) and text[i + 1] in {'.', '\uff0e', '\u3002'}
                if prev_is_dot or next_is_dot:
                    continue
            if cls._is_after_digit_or_latin(text, i):
                continue
            chars[i] = replacements[ch]
        return ''.join(chars)

    @classmethod
    def _normalize_bracket_pairs(cls, text):
        pair_sets = [
            ({'(': '\uff08', '\uff08': '\uff08'}, {')': '\uff09', '\uff09': '\uff09'}),
            ({'[': '\uff3b', '\uff3b': '\uff3b'}, {']': '\uff3d', '\uff3d': '\uff3d'}),
        ]
        result = text

        for open_chars, close_chars in pair_sets:
            chars = list(result)
            stack = []
            for i, ch in enumerate(chars):
                if ch in open_chars:
                    stack.append(i)
                elif ch in close_chars and stack:
                    open_index = stack.pop()
                    content = ''.join(chars[open_index + 1:i])
                    if not cls._has_chinese(content):
                        continue
                    if cls._is_after_digit_or_latin(result, open_index):
                        continue
                    if cls._is_before_digit_or_latin(result, i):
                        continue
                    chars[open_index] = open_chars[chars[open_index]]
                    chars[i] = close_chars[ch]
            result = ''.join(chars)

        return result

    @classmethod
    def _normalize_quote_pairs(cls, text, quote_chars, left_quote, right_quote, skip_inner_latin=False):
        chars = list(text)
        open_index = None

        for i, ch in enumerate(chars):
            if ch not in quote_chars:
                continue

            prev_is_latin = cls._is_after_digit_or_latin(text, i)
            next_is_latin = cls._is_before_digit_or_latin(text, i)
            if skip_inner_latin and prev_is_latin and next_is_latin:
                continue

            if open_index is None:
                if prev_is_latin:
                    continue
                open_index = i
                continue

            content = ''.join(chars[open_index + 1:i])
            should_normalize = (
                cls._has_chinese(content)
                and not cls._is_after_digit_or_latin(text, open_index)
                and not cls._is_before_digit_or_latin(text, i)
            )
            if should_normalize:
                chars[open_index] = left_quote
                chars[i] = right_quote
            open_index = None

        return ''.join(chars)

    @classmethod
    def _normalize_symbols_in_text(cls, text):
        if not text or not cls._has_chinese(text):
            return text

        result = cls._normalize_bracket_pairs(text)
        result = cls._normalize_ellipsis(result)
        result = cls._normalize_quote_pairs(
            result,
            {'"', '\u201c', '\u201d', '\u201e', '\u201f', '\u300c', '\u300d'},
            '\u201c',
            '\u201d',
        )
        result = cls._normalize_quote_pairs(
            result,
            {"'", '\u2018', '\u2019', '\u201a', '\u201b'},
            '\u2018',
            '\u2019',
            skip_inner_latin=True,
        )
        result = cls._normalize_simple_punctuation(result)
        return result

    @staticmethod
    def _redistribute_text_to_runs(runs, new_full_text):
        if not runs:
            return

        run_lengths = [len(run.text) for run in runs]
        if len(new_full_text) == sum(run_lengths):
            pos = 0
            for run, length in zip(runs, run_lengths):
                run.text = new_full_text[pos:pos + length]
                pos += length
            return

        runs[0].text = new_full_text
        for run in runs[1:]:
            run.text = ''

    def _normalize_paragraph_symbols(self, para):
        text = para.text
        if not text.strip() or not para.runs:
            return False
        if '<w:fldChar' in para._p.xml or '<w:instrText' in para._p.xml:
            return False

        normalized = self._normalize_symbols_in_text(text)
        if normalized == text:
            return False

        self._redistribute_text_to_runs(para.runs, normalized)
        return True

    def _normalize_document_symbols(self, doc):
        changes = 0

        for para in doc.paragraphs:
            if self._normalize_paragraph_symbols(para):
                changes += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if self._normalize_paragraph_symbols(para):
                            changes += 1

        return changes

    # ------------------------------------------------------------------
    # Markdown cleanup
    # ------------------------------------------------------------------
    @staticmethod
    def _clean_markdown(text):
        """
        Clean Markdown content to plain text:
        1. Remove images, links, HTML, inline code markers
        2. Remove bold/italic markers (*, **, __)
        3. Remove heading markers (#)
        4. Remove blockquote markers (>)
        5. Remove horizontal rules (---)
        6. Fix auto-numbering (1. 1. 1.) to sequential (1. 2. 3.)
        """
        if not text:
            return ""

        # Global inline element replacements
        # Remove images: ![alt](url) -> alt
        text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
        # Remove links: [text](url) -> text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Remove inline code: `code` -> code
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove bold/italic (** or __)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)

        lines = text.split('\n')
        new_lines = []
        list_counters = {}  # indent -> count
        in_list_block = False

        for line in lines:
            cleaned_line = line

            # Remove heading markers: # Title -> Title
            header_match = re.match(r'^\s*#+\s+(.*)', cleaned_line)
            if header_match:
                cleaned_line = header_match.group(1)
                in_list_block = False
                list_counters = {}

            # Remove blockquote markers: > Text -> Text
            blockquote_match = re.match(r'^\s*>\s?(.*)', cleaned_line)
            if blockquote_match:
                cleaned_line = blockquote_match.group(1)

            # Remove horizontal rules: ---, ***, ___
            if re.match(r'^\s*[-*_]{3,}\s*$', cleaned_line):
                cleaned_line = ""

            # List processing (auto-numbering fix)
            list_match = re.match(r'^(\s*)(\d+)\.(\s+.*)', cleaned_line)
            if list_match:
                indent_str = list_match.group(1)
                indent = len(indent_str.expandtabs(4))
                content = list_match.group(3)

                if indent not in list_counters:
                    list_counters[indent] = 0

                if not in_list_block and int(list_match.group(2)) == 1:
                    list_counters[indent] = 0

                list_counters[indent] += 1
                cleaned_line = f"{indent_str}{list_counters[indent]}.{content}"
                in_list_block = True
            else:
                if cleaned_line.strip() != '':
                    if re.match(r'^\s*[*+-]\s+', cleaned_line):
                        in_list_block = False
                        list_counters = {}

            # Remove remaining * italic markers (careful not to break list markers)
            is_bullet = re.match(r'^\s*[*+-]\s', cleaned_line)

            def remove_stars(m):
                return m.group(1)

            if is_bullet:
                bullet_match = re.match(r'^(\s*[*+-]\s)(.*)', cleaned_line)
                if bullet_match:
                    marker = bullet_match.group(1)
                    content = bullet_match.group(2)
                    content = re.sub(r'(?<!\\)\*([^\s*][^*]*?)(?<!\\)\*', remove_stars, content)
                    cleaned_line = marker + content
            else:
                cleaned_line = re.sub(r'(?<!\\)\*([^\s*][^*]*?)(?<!\\)\*', remove_stars, cleaned_line)

            new_lines.append(cleaned_line)

        return '\n'.join(new_lines)

    # ------------------------------------------------------------------
    # Blank line removal for plain text sources
    # ------------------------------------------------------------------
    @staticmethod
    def _remove_blank_lines_from_text(text):
        """
        Remove extra blank lines from txt/md text:
        - Single blank line: delete
        - 2+ consecutive blank lines: merge to 1
        """
        if not text:
            return text

        lines = text.split('\n')
        result = []
        blank_count = 0

        for line in lines:
            if line.strip() == '':
                blank_count += 1
            else:
                if blank_count >= 2:
                    result.append('')
                result.append(line)
                blank_count = 0

        if blank_count >= 2:
            result.append('')

        return '\n'.join(result)

    # ------------------------------------------------------------------
    # Text file reading
    # ------------------------------------------------------------------
    @staticmethod
    def _read_text_file(path):
        """Try multiple encodings to read text file"""
        for enc in ['utf-8', 'utf-8-sig', 'gbk', 'gb18030']:
            try:
                with open(path, 'r', encoding=enc) as f:
                    content = f.read()
                return content
            except UnicodeDecodeError:
                continue
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def convert_to_docx(self, input_path):
        file_ext = os.path.splitext(input_path)[1].lower()
        is_from_txt = (file_ext in ('.txt', '.md'))
        temp_dir = os.path.dirname(input_path)
        base_name = os.path.splitext(os.path.basename(input_path))[0]

        if file_ext == '.docx':
            self._log("检测到 .docx 文件，正在创建安全的处理副本...")
            temp_docx_path = os.path.join(temp_dir, f"~temp_copy_{base_name}.docx")
            shutil.copy2(input_path, temp_docx_path)
            self.temp_files.append(temp_docx_path)
            self._log(f"  > 副本创建成功: {os.path.basename(temp_docx_path)}")
            return temp_docx_path, False

        temp_docx_path = os.path.join(temp_dir, f"~temp_converted_{base_name}.docx")
        self.temp_files.append(temp_docx_path)

        if file_ext == '.txt':
            self._log("检测到 .txt 文件，正在创建 .docx...")
            text_content = self._read_text_file(input_path)
            if self.remove_blank_lines:
                text_content = self._remove_blank_lines_from_text(text_content)
                self._log("  > 已删除 TXT 中的多余空行。")
            doc = Document()
            for line in text_content.split('\n'):
                doc.add_paragraph(line.strip())
            doc.save(temp_docx_path)
            self._log("TXT转换完成。")
            return temp_docx_path, is_from_txt
        elif file_ext == '.md':
            self._log("检测到 .md 文件，正在清理 Markdown 标记并创建 .docx...")
            raw_text = self._read_text_file(input_path)
            cleaned_text = self._clean_markdown(raw_text)
            if self.remove_blank_lines:
                cleaned_text = self._remove_blank_lines_from_text(cleaned_text)
                self._log("  > 已删除 Markdown 文本中的多余空行。")
            doc = Document()
            for line in cleaned_text.split('\n'):
                doc.add_paragraph(line.strip())
            doc.save(temp_docx_path)
            self._log("Markdown 转换完成。")
            return temp_docx_path, is_from_txt
        elif file_ext in ['.wps', '.doc']:
            self._log(f"正在转换 {file_ext} 文件为 .docx...")
            app = self._get_wps_app()
            doc_com = app.Documents.Open(os.path.abspath(input_path), ReadOnly=1)
            doc_com.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=12)
            doc_com.Close()
            self._log("文件格式转换完成。")
            return temp_docx_path, is_from_txt
        
        raise ValueError(f"不支持的文件格式: {file_ext}")

    def _preprocess_com_tasks(self, docx_path):
        self._log("正在对副本执行预处理（接受所有修订、转换自动编号）...")
        app = self._get_wps_app()
        try:
            doc_com = app.Documents.Open(os.path.abspath(docx_path))
            
            doc_com.TrackRevisions = False
            self._log("  > 已关闭修订追踪。")
            
            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions()
                self._log("  > 已接受文档副本中的所有修订。")
            
            doc_com.Content.ListFormat.ConvertNumbersToText()
            self._log("  > 已将副本中的自动编号转换为文本。")
            
            if doc_com.Revisions.Count > 0:
                doc_com.AcceptAllRevisions()
                self._log("  > 已接受编号转换产生的修订。")
            
            doc_com.TrackRevisions = False
            
            doc_com.Save()
            doc_com.Close()
            self._log("预处理完成。")
        except Exception as e:
            self._log(f"警告：执行预处理任务时出错: {e}")

    def _create_page_number(self, paragraph, text):
        font_name = self.config['page_number_font']
        font_size = self.config['page_number_size']
        self._set_run_font(paragraph.add_run('— '), font_name, font_size, set_color=True)
        run_field = paragraph.add_run()
        self._set_run_font(run_field, font_name, font_size, set_color=True)
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = text
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run_field._r.extend([fldChar1, instrText, fldChar2])
        self._set_run_font(paragraph.add_run(' —'), font_name, font_size, set_color=True)

    def _apply_page_setup(self, doc, is_from_txt=False):
        self._log("正在应用页面边距和页码设置...")
        
        # 判断是否需要强制设置A4纸
        # 逻辑：如果是纯文本来源（包括直接输入）或者 用户勾选了强制A4，则设置为A4
        should_set_a4 = is_from_txt or self.config.get('force_a4', False)

        for section in doc.sections:
            section.top_margin = Cm(self.config['margin_top'])
            section.bottom_margin = Cm(self.config['margin_bottom'])
            section.left_margin = Cm(self.config['margin_left'])
            section.right_margin = Cm(self.config['margin_right'])
            section.footer_distance = Cm(self.config['footer_distance'])

            # 设置纸张大小为A4 (仅在需要时)
            if should_set_a4:
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)

            if self.config['page_number_align'] == '居中':
                p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
                p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; self._create_page_number(p, 'PAGE')
            elif self.config['page_number_align'] == '奇偶分页':
                doc.settings.odd_and_even_pages_header_footer = True
                footer_odd = section.footer
                p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
                p_odd.clear(); p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT; self._create_page_number(p_odd, 'PAGE')
                
                footer_even = section.even_page_footer
                p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
                p_even.clear(); p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT; self._create_page_number(p_even, 'PAGE')
        
        if should_set_a4:
            self._log("  > 已将页面大小设置为 A4。")

    def _set_run_font(self, run, font_name, size_pt, set_color=False):
        run.font.size = Pt(size_pt)
        if set_color: run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for theme_attr in ('w:eastAsiaTheme', 'w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme', 'w:csTheme'):
            rFonts.attrib.pop(qn(theme_attr), None)
        rFonts.set(qn('w:eastAsia'), font_name)
        # 根据配置决定西文字体（数字、字母）
        en_font = self.config.get('english_font') if self.config.get('use_custom_english_font', False) else font_name
        en_font = en_font or font_name
        run.font.name = en_font
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)

    def _apply_font_to_runs(self, para, font_name, size_pt, set_color=False):
        for run in para.runs: self._set_run_font(run, font_name, size_pt, set_color=set_color)

    def _get_paragraph_font_info(self, para):
        """获取段落主要字体和字号信息"""
        if not para.runs:
            return None, None
        
        # 获取第一个非空run的字体信息
        for run in para.runs:
            if run.text.strip():
                font_name = run.font.name
                font_size = run.font.size.pt if run.font.size else None
                return font_name, font_size
        return None, None

    def _strip_leading_whitespace(self, para):
        if not para.runs: return
        while para.runs and not para.runs[0].text.strip():
            p = para._p
            p.remove(para.runs[0]._r)
        if not para.runs: return
        first_run = para.runs[0]
        original_text = first_run.text
        stripped_text = original_text.lstrip()
        if original_text != stripped_text:
            first_run.text = stripped_text
            self._log("  > 已移除段落前的多余空格。")
    
    def _reset_pagination_properties(self, para):
        para.paragraph_format.widow_control = False
        para.paragraph_format.keep_with_next = False
        para.paragraph_format.keep_lines_together = False
        para.paragraph_format.page_break_before = False
        para.paragraph_format.keep_together = False

    def _get_outline_level(self, para):
        """
        读取段落的当前大纲级别
        返回: 0-8 表示级别1-9，None 表示未设置
        """
        pPr = para._p.get_or_add_pPr()
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is not None:
            val = outlineLvl.get(qn('w:val'))
            if val is not None:
                return int(val)
        return None

    def _set_outline_level(self, para, level):
        """
        直接设置段落的大纲级别，不通过样式，不影响字体字号等格式
        level: 1-9 的整数，表示大纲级别
        返回: 原有的大纲级别 (0-8) 或 None
        """
        if level < 1 or level > 9:
            self._log(f"  > 警告：大纲级别 {level} 超出范围 (1-9)，已跳过设置")
            return None
        
        # 读取原有大纲级别
        original_level = self._get_outline_level(para)
        
        # 设置新的大纲级别 (Word内部用0-8表示1-9级)
        pPr = para._p.get_or_add_pPr()
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), str(level - 1))
        
        return original_level

    def _format_heading(self, para, level):
        """
        为段落设置大纲级别（仅设置大纲级别，不影响其他格式）
        """
        if not self.config['set_outline']:
            self._log(f"  > 大纲级别设置已禁用，跳过")
            return
        
        # 获取段落文本预览用于日志
        text_preview = para.text.strip()[:30].replace("\n", " ")
        
        original_level = self._set_outline_level(para, level)
        
        if original_level is not None:
            self._log(f"  > 大纲级别: Lv{original_level + 1} → Lv{level} (覆盖) - \"{text_preview}...\"")
        else:
            self._log(f"  > 大纲级别: 无 → Lv{level} (新设) - \"{text_preview}...\"")

    def _apply_text_indent_and_align(self, para):
        pf = para.paragraph_format
        # 清除 python-docx 层面的缩进
        pf.first_line_indent = None
        pf.left_indent = Cm(self.config['left_indent_cm'])
        pf.right_indent = Cm(self.config['right_indent_cm'])
        
        # 操作底层 XML，彻底清理残留的缩进属性，避免与首行缩进叠加
        ind = para._p.get_or_add_pPr().get_or_add_ind()
        # 清除可能残留的字符单位左缩进（防止与首行缩进叠加显示为4字符）
        ind.attrib.pop(qn('w:leftChars'), None)
        # 清除可能残留的悬挂缩进
        ind.attrib.pop(qn('w:hanging'), None)
        ind.attrib.pop(qn('w:hangingChars'), None)
        # 清除可能残留的固定值首行缩进（我们使用字符单位 firstLineChars）
        ind.attrib.pop(qn('w:firstLine'), None)
        # 设置首行缩进 2 字符（200 = 2 × 100）
        ind.set(qn("w:firstLineChars"), "200")
        
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _iter_block_items(self, parent):
        parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P): yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl): yield Table(child, parent)

    def _get_or_add_table_pr(self, table):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        return tbl_pr

    def _set_table_borders(self, table, size_pt=0.5, color="000000"):
        size = max(1, int(float(size_pt) * 8))
        tbl_pr = self._get_or_add_table_pr(table)
        borders = tbl_pr.find(qn('w:tblBorders'))
        if borders is None:
            borders = OxmlElement('w:tblBorders')
            tbl_pr.append(borders)
        else:
            for child in list(borders):
                borders.remove(child)

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(size))
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), color)
            borders.append(elem)

    def _set_cell_borders(self, cell, size_pt=0.5, color="000000"):
        size = max(1, int(float(size_pt) * 8))
        tc = cell._tc
        tc_pr = tc.tcPr
        if tc_pr is None:
            tc_pr = OxmlElement('w:tcPr')
            tc.insert(0, tc_pr)

        borders = tc_pr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        else:
            for child in list(borders):
                borders.remove(child)

        for edge in ('top', 'left', 'bottom', 'right'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(size))
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), color)
            borders.append(elem)

    def _set_table_cell_margins(self, table, top_cm=0.0, bottom_cm=0.0, left_cm=0.05, right_cm=0.05):
        tbl_pr = self._get_or_add_table_pr(table)
        cell_mar = tbl_pr.find(qn('w:tblCellMar'))
        if cell_mar is None:
            cell_mar = OxmlElement('w:tblCellMar')
            tbl_pr.append(cell_mar)

        def set_side(tag, cm_value):
            node = cell_mar.find(qn(f'w:{tag}'))
            if node is None:
                node = OxmlElement(f'w:{tag}')
                cell_mar.append(node)
            node.set(qn('w:type'), 'dxa')
            node.set(qn('w:w'), str(int(Cm(float(cm_value)).twips)))

        set_side('top', top_cm)
        set_side('bottom', bottom_cm)
        set_side('left', left_cm)
        set_side('right', right_cm)

    def _set_table_width_percent(self, table, percent=100):
        percent = max(1, min(100, int(float(percent))))
        tbl_pr = self._get_or_add_table_pr(table)
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None:
            tbl_w = OxmlElement('w:tblW')
            tbl_pr.append(tbl_w)
        tbl_w.set(qn('w:type'), 'pct')
        tbl_w.set(qn('w:w'), str(percent * 50))

    def _set_table_indent(self, table, indent_twips=0):
        tbl_pr = self._get_or_add_table_pr(table)
        tbl_ind = tbl_pr.find(qn('w:tblInd'))
        if tbl_ind is None:
            tbl_ind = OxmlElement('w:tblInd')
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn('w:type'), 'dxa')
        tbl_ind.set(qn('w:w'), str(int(indent_twips)))

    @staticmethod
    def _table_text_weight(text):
        weight = 0.0
        for ch in text:
            weight += 0.5 if ord(ch) < 128 else 1.0
        return weight

    @staticmethod
    def _normalize_table_pcts(weights, min_pct, max_pct):
        total = sum(weights) or 1.0
        pcts = [w / total * 100 for w in weights]
        for i, value in enumerate(pcts):
            if value < min_pct:
                pcts[i] = min_pct
            elif value > max_pct:
                pcts[i] = max_pct
        total = sum(pcts) or 1.0
        return [value / total * 100 for value in pcts]

    def _set_table_col_widths_by_content(self, table, min_pct=8, max_pct=45):
        if not table.rows:
            return
        col_count = max(len(row.cells) for row in table.rows)
        if col_count == 0:
            return

        min_pct = max(1.0, float(min_pct))
        max_pct = max(min_pct, float(max_pct))
        max_weights = [1.0] * col_count
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                text = ''.join(p.text for p in cell.paragraphs).strip()
                if text:
                    max_weights[col_idx] = max(max_weights[col_idx], self._table_text_weight(text))

        pcts = self._normalize_table_pcts(max_weights, min_pct, max_pct)
        tbl = table._tbl
        tbl_grid = tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement('w:tblGrid')
            tbl.insert(0, tbl_grid)
        else:
            for child in list(tbl_grid):
                tbl_grid.remove(child)

        for pct in pcts:
            grid_col = OxmlElement('w:gridCol')
            grid_col.set(qn('w:w'), str(int(pct * 50)))
            tbl_grid.append(grid_col)

        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tc_pr = tc.tcPr
                if tc_pr is None:
                    tc_pr = OxmlElement('w:tcPr')
                    tc.insert(0, tc_pr)
                tc_w = tc_pr.find(qn('w:tcW'))
                if tc_w is None:
                    tc_w = OxmlElement('w:tcW')
                    tc_pr.append(tc_w)
                tc_w.set(qn('w:type'), 'pct')
                tc_w.set(qn('w:w'), str(int(pcts[col_idx] * 50)))

    @staticmethod
    def _is_numeric_table_text(text):
        text = (text or '').strip()
        if not text:
            return False
        text = text.replace(',', '').replace('，', '').replace('％', '%')
        text = re.sub(r'^[¥￥$]', '', text)
        text = re.sub(r'(元|万元|亿元)$', '', text)
        return re.match(r'^[-+]?(?:\d+(?:\.\d+)?|\.\d+)%?$', text) is not None

    @staticmethod
    def _is_short_table_text(text, max_len=4):
        text = (text or '').strip()
        return 0 < len(text) <= int(max_len)

    @staticmethod
    def _config_float(config, key, default):
        value = config.get(key, default)
        try:
            if value == '':
                return default
            return float(value)
        except (TypeError, ValueError):
            return default

    def _format_tables(self, doc, apply_color=True):
        if not self.config.get('enable_table_formatting', False):
            self._log("表格自动调整未启用，跳过表格内容格式化。")
            return

        tables = list(doc.tables)
        if not tables:
            self._log("未发现表格，跳过表格内容格式化。")
            return

        table_font = self.config.get('table_font', self.config.get('body_font', '仿宋_GB2312'))
        table_header_font = self.config.get('table_header_font', table_font)
        table_size = self._config_float(self.config, 'table_size', self.config.get('body_size', 12))
        table_line_spacing = self._config_float(self.config, 'table_line_spacing', 22)
        row_height_cm = self._config_float(self.config, 'table_row_height_cm', 0.7)
        border_size_pt = self._config_float(self.config, 'table_border_size_pt', 0.5)
        width_percent = self._config_float(self.config, 'table_width_percent', 100)
        col_min_pct = self._config_float(self.config, 'table_col_min_pct', 8)
        col_max_pct = self._config_float(self.config, 'table_col_max_pct', 45)
        short_text_len = self._config_float(self.config, 'table_short_text_len', 4)
        auto_col_width = self.config.get('table_auto_col_width', True)
        header_bold = self.config.get('table_header_bold', True)
        smart_align = self.config.get('table_smart_align', False)
        unified_borders = self.config.get('table_unified_borders', True)

        self._log(f"开始格式化表格内容（共 {len(tables)} 个）...")
        for table_idx, table in enumerate(tables, start=1):
            self._log(f"  > 表格 {table_idx}: 调整宽度、行高、字体和单元格格式")
            table.autofit = not auto_col_width
            self._set_table_width_percent(table, width_percent)
            self._set_table_indent(table, 0)
            self._set_table_cell_margins(table)
            if unified_borders:
                self._set_table_borders(table, size_pt=border_size_pt)
            if auto_col_width:
                self._set_table_col_widths_by_content(table, min_pct=col_min_pct, max_pct=col_max_pct)

            serial_col_idx = None
            if table.rows:
                for col_idx, cell in enumerate(table.rows[0].cells):
                    head_text = ''.join(p.text for p in cell.paragraphs).strip()
                    if '序号' in head_text or head_text == '序':
                        serial_col_idx = col_idx
                        break

            for row_idx, row in enumerate(table.rows):
                if row_height_cm > 0:
                    row.height = Cm(row_height_cm)
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

                for col_idx, cell in enumerate(row.cells):
                    if unified_borders:
                        self._set_cell_borders(cell, size_pt=border_size_pt)

                    cell_text = ''.join(p.text for p in cell.paragraphs).strip()
                    for para in cell.paragraphs:
                        if para.text.strip():
                            for run in para.runs:
                                font_name = table_header_font if row_idx == 0 else table_font
                                self._set_run_font(run, font_name, table_size, set_color=apply_color)
                                if row_idx == 0 and header_bold:
                                    run.font.bold = True

                        para.paragraph_format.first_line_indent = Pt(0)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        if table_line_spacing > 0:
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                            para.paragraph_format.line_spacing = Pt(table_line_spacing)
                        else:
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                        if smart_align:
                            if row_idx == 0:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif '合计' in cell_text or '总计' in cell_text:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif serial_col_idx is not None and col_idx == serial_col_idx:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif self._is_numeric_table_text(cell_text):
                                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif self._is_short_table_text(cell_text, short_text_len):
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _find_title_and_subtitle_paragraphs(self, doc, is_from_txt, start_index=0):
        """
        查找题目和副标题段落的索引范围
        返回: (title_indices, subtitle_indices)
        title_indices: 题目行的索引列表
        subtitle_indices: 副标题行的索引列表
        """
        ch_num = r'[一二三四五六七八九十百千万零]+'
        re_h1 = re.compile(r'^' + ch_num + r'\s*、')
        re_h2 = re.compile(r'^[（\(]' + ch_num + r'[）\)]')

        all_blocks = list(self._iter_block_items(doc))
        
        # 查找首个标题行
        first_title_idx = -1
        
        if is_from_txt:
            self._log("文档源自 TXT，采用智能规则查找题目...")
            for idx in range(start_index, len(all_blocks)):
                block = all_blocks[idx]
                if isinstance(block, Paragraph) and block.text.strip():
                    text_to_check = block.text.strip()
                    if re_h1.match(text_to_check) or re_h2.match(text_to_check):
                        self._log(f"  > 首个非空行 (块 {idx + 1}) 符合标题格式，认定本文档无独立题目。")
                        return [], []
                    else:
                        self._log(f"  > 在块 {idx + 1} 发现首个非空段落，认定为题目首行。")
                        first_title_idx = idx
                        break
        else:
            self._log("正在预扫描以确定居中题目位置...")
            for idx in range(start_index, len(all_blocks)):
                block = all_blocks[idx]
                if not isinstance(block, Paragraph) or not block.text.strip(): 
                    continue
                para = block
                text_to_check = para.text.lstrip()
                if re_h1.match(text_to_check) or re_h2.match(text_to_check):
                    self._log("  > 发现一级/二级标题，在此之前未找到居中题目。")
                    return [], []
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    self._log(f"  > 在块 {idx + 1} 发现潜在题目首行。")
                    first_title_idx = idx
                    break
        
        if first_title_idx == -1:
            self._log("  > 扫描结束，未能找到题目。")
            return [], []
        
        # 获取首个标题行的字体字号信息
        first_title_para = all_blocks[first_title_idx]
        title_font, title_size = self._get_paragraph_font_info(first_title_para)
        
        # 向下查找连续的标题行
        title_indices = [first_title_idx]
        idx = first_title_idx + 1
        
        while idx < len(all_blocks):
            block = all_blocks[idx]
            if not isinstance(block, Paragraph):
                break
            
            para = block
            text = para.text.strip()
            
            # 遇到空行，停止标题识别
            if not text:
                self._log(f"  > 在块 {idx + 1} 遇到空行，标题识别结束。")
                break
            
            # 检查是否居中
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                break
            
            # 检查字体字号是否与首行相同
            para_font, para_size = self._get_paragraph_font_info(para)
            if para_font == title_font and para_size == title_size:
                self._log(f"  > 块 {idx + 1} 也是标题行（居中且字体字号相同）。")
                title_indices.append(idx)
                idx += 1
            else:
                # 字体字号不同，可能是副标题的开始
                break
        
        self._log(f"  > 共识别到 {len(title_indices)} 行标题。")
        
        # 查找副标题
        subtitle_indices = []
        subtitle_start_idx = idx
        
        # 跳过空行
        while subtitle_start_idx < len(all_blocks):
            block = all_blocks[subtitle_start_idx]
            if isinstance(block, Paragraph) and block.text.strip():
                break
            if isinstance(block, Paragraph):
                subtitle_start_idx += 1
            else:
                # 遇到非段落（如表格），停止
                break
        
        # 检查是否有副标题
        if subtitle_start_idx < len(all_blocks):
            block = all_blocks[subtitle_start_idx]
            if isinstance(block, Paragraph):
                para = block
                text = para.text.strip()
                
                # 副标题必须居中
                if text and para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    # 检查字体字号是否与标题不同
                    para_font, para_size = self._get_paragraph_font_info(para)
                    if para_font != title_font or para_size != title_size:
                        self._log(f"  > 在块 {subtitle_start_idx + 1} 发现副标题首行（居中且字体字号与标题不同）。")
                        subtitle_indices.append(subtitle_start_idx)
                        
                        # 查找连续的副标题行
                        subtitle_font, subtitle_size = para_font, para_size
                        idx = subtitle_start_idx + 1
                        
                        while idx < len(all_blocks):
                            block = all_blocks[idx]
                            if not isinstance(block, Paragraph):
                                break
                            
                            para = block
                            text = para.text.strip()
                            
                            # 遇到空行，停止副标题识别
                            if not text:
                                self._log(f"  > 在块 {idx + 1} 遇到空行，副标题识别结束。")
                                break
                            
                            # 检查是否居中
                            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                                break
                            
                            # 检查字体字号是否与副标题首行相同
                            para_font, para_size = self._get_paragraph_font_info(para)
                            if para_font == subtitle_font and para_size == subtitle_size:
                                self._log(f"  > 块 {idx + 1} 也是副标题行（居中且字体字号相同）。")
                                subtitle_indices.append(idx)
                                idx += 1
                            else:
                                break
                        
                        self._log(f"  > 共识别到 {len(subtitle_indices)} 行副标题。")
        
        return title_indices, subtitle_indices

    def format_document(self, input_path, output_path):
        processing_path, is_from_txt = self.convert_to_docx(input_path)
        if not is_from_txt: self._preprocess_com_tasks(processing_path)
        
        doc = Document(processing_path)

        if self.config.get('normalize_punctuation', False):
            symbol_changes = self._normalize_document_symbols(doc)
            self._log(f"符号标准化完成，共修复 {symbol_changes} 个段落/表格单元格。")
        
        all_blocks = list(self._iter_block_items(doc))
        processed_indices = set()
        
        apply_color = not is_from_txt

        if not is_from_txt:
            self._log("正在扫描图表标题...")
            for idx, block in enumerate(all_blocks):
                is_pic_para = isinstance(block, Paragraph) and ('<w:drawing>' in block._p.xml or '<w:pict>' in block._p.xml)
                is_table = isinstance(block, Table)
                
                if not (is_pic_para or is_table): continue
                
                for direction in [-1, 1]:
                    caption_found = False
                    for i in range(idx + direction, -1 if direction == -1 else len(all_blocks), direction):
                        if i in processed_indices: continue
                        potential_caption = all_blocks[i]
                        if not isinstance(potential_caption, Paragraph): break 
                        text = potential_caption.text.strip()
                        if text: 
                            if potential_caption.alignment == WD_ALIGN_PARAGRAPH.CENTER and (text.startswith("图") or text.startswith("表")):
                                detected_type = "图" if text.startswith("图") else "表"
                                self._log(f"  > 发现 {detected_type} 的标题: \"{text[:30]}...\" (在段落 {i+1})")
                                config_font_key = f'{("figure" if detected_type == "图" else "table")}_caption_font'
                                config_size_key = f'{("figure" if detected_type == "图" else "table")}_caption_size'
                                config_font = self.config[config_font_key]
                                config_size = self.config[config_size_key]
                                self._apply_font_to_runs(potential_caption, config_font, config_size, set_color=apply_color)
                                processed_indices.add(i)
                                caption_found = True
                            break 
                    if caption_found: break 

        # 查找主标题和副标题
        title_indices, subtitle_indices = self._find_title_and_subtitle_paragraphs(doc, is_from_txt)
        
        # 将标题和副标题索引加入已处理集合
        for idx in title_indices:
            processed_indices.add(idx)
        for idx in subtitle_indices:
            processed_indices.add(idx)

        self._log("预扫描完成，开始逐段格式化...")
        if self.config['set_outline']:
            self._log("【大纲级别设置已启用】")
        else:
            self._log("【大纲级别设置已禁用】")
            
        re_h1 = re.compile(r'^[一二三四五六七八九十百千万零]+\s*、')
        re_h2 = re.compile(r'^[（\(][一二三四五六七八九十百千万零]+[）\)]')
        re_h3 = re.compile(r'^\d+\s*[\.．]')
        re_h4 = re.compile(r'^[（\(]\d+[）\)]')
        re_attachment = re.compile(r'^附件\s*(\d+|[一二三四五六七八九十百千万零]+)?\s*[:：]?$')

        # 格式化主标题
        if title_indices:
            self._log(f"\n开始格式化主标题（共 {len(title_indices)} 行）...")
            for idx in title_indices:
                para = all_blocks[idx]
                self._log(f"段落 {idx + 1}: 主标题行 - \"{para.text[:30]}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['title_font'], self.config['title_size'], set_color=apply_color)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = None
                
                # 设置标题行间距
                spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                spacing.set(qn('w:beforeAutospacing'), '0')
                spacing.set(qn('w:afterAutospacing'), '0')
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = Pt(self.config['title_line_spacing'])
                
                self._reset_pagination_properties(para)
        
        # 格式化副标题
        if subtitle_indices:
            self._log(f"\n开始格式化副标题（共 {len(subtitle_indices)} 行）...")
            for idx in subtitle_indices:
                para = all_blocks[idx]
                self._log(f"段落 {idx + 1}: 副标题行 - \"{para.text[:30]}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['subtitle_font'], self.config['subtitle_size'], set_color=apply_color)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = None
                
                # 设置副标题行间距
                spacing = para._p.get_or_add_pPr().get_or_add_spacing()
                spacing.set(qn('w:beforeAutospacing'), '0')
                spacing.set(qn('w:afterAutospacing'), '0')
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = Pt(self.config['subtitle_line_spacing'])
                
                self._reset_pagination_properties(para)

        block_idx = 0
        while block_idx < len(all_blocks):
            block = all_blocks[block_idx]
            
            if block_idx in processed_indices:
                if block_idx not in title_indices and block_idx not in subtitle_indices:
                    self._log(f"块 {block_idx + 1}: 已作为图表/附件标题处理 - 跳过")
                block_idx += 1
                continue

            current_block_num = block_idx + 1
            if isinstance(block, Table): 
                self._log(f"块 {current_block_num}: 表格 - 跳过"); block_idx += 1; continue
            
            para = block
            if not para.text.strip(): 
                self._log(f"段落 {current_block_num}: 空白 - 跳过"); block_idx += 1; continue
            
            is_pic = '<w:drawing>' in para._p.xml or '<w:pict>' in para._p.xml
            is_embedded_obj = '<w:object>' in para._p.xml
            if is_pic or is_embedded_obj:
                log_msg = "图片" if is_pic else "附件"
                self._log(f"段落 {current_block_num}: {log_msg} - 仅格式化文字")
                
                text_to_check = para.text.lstrip()
                para_text_preview = text_to_check[:30].replace("\n", " ")

                if re_h1.match(text_to_check):
                    self._log(f"  > 文字识别为一级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color)
                elif re_h2.match(text_to_check):
                    self._log(f"  > 文字识别为二级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                elif re_h3.match(text_to_check):
                    self._log(f"  > 文字识别为三级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                elif re_h4.match(text_to_check):
                    self._log(f"  > 文字识别为四级标题: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                elif text_to_check:
                    self._log(f"  > 文字识别为正文: \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)

                block_idx += 1
                continue

            original_text, text_to_check = para.text, para.text.lstrip()
            text_to_check_stripped = para.text.strip()
            leading_space_count = len(original_text) - len(text_to_check)
            para_text_preview = text_to_check[:30].replace("\n", " ")
            
            spacing = para._p.get_or_add_pPr().get_or_add_spacing()
            spacing.set(qn('w:beforeAutospacing'), '0'); spacing.set(qn('w:afterAutospacing'), '0')
            para.paragraph_format.space_before, para.paragraph_format.space_after = Pt(0), Pt(0)
            para.paragraph_format.line_spacing = Pt(self.config['line_spacing'])

            is_attachment_enabled = self.config.get('enable_attachment_formatting', False)
            is_attachment_candidate = False
            if is_from_txt:
                if re_attachment.match(text_to_check_stripped): is_attachment_candidate = True
            elif para.alignment in [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY, None] and re_attachment.match(text_to_check_stripped):
                is_attachment_candidate = True

            if is_attachment_enabled and is_attachment_candidate:
                self._log(f"段落 {current_block_num}: 附件标识 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['attachment_font'], self.config['attachment_size'], set_color=apply_color)
                self._reset_pagination_properties(para)
                para.paragraph_format.page_break_before = True
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = None
                
                ind = para._p.get_or_add_pPr().get_or_add_ind()
                ind.set(qn("w:firstLineChars"), "0")
                
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._format_heading(para, 1)

                # 查找并格式化附件的标题和副标题
                search_idx = block_idx + 1
                
                # 查找附件的标题和副标题
                att_title_indices, att_subtitle_indices = self._find_title_and_subtitle_paragraphs(doc, is_from_txt, search_idx)
                
                # 将附件的标题和副标题加入已处理集合
                for idx in att_title_indices:
                    processed_indices.add(idx)
                for idx in att_subtitle_indices:
                    processed_indices.add(idx)
                
                # 格式化附件的标题
                if att_title_indices:
                    self._log(f"  > 识别到附件标题（共 {len(att_title_indices)} 行）")
                    for idx in att_title_indices:
                        para_title = all_blocks[idx]
                        self._log(f"    段落 {idx + 1}: 附件标题行 - \"{para_title.text.strip()[:30]}...\"")
                        self._strip_leading_whitespace(para_title)
                        self._apply_font_to_runs(para_title, self.config['title_font'], self.config['title_size'], set_color=apply_color)
                        para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para_title.paragraph_format.first_line_indent = None
                        
                        # 设置标题行间距
                        spacing = para_title._p.get_or_add_pPr().get_or_add_spacing()
                        spacing.set(qn('w:beforeAutospacing'), '0')
                        spacing.set(qn('w:afterAutospacing'), '0')
                        para_title.paragraph_format.space_before = Pt(0)
                        para_title.paragraph_format.space_after = Pt(0)
                        para_title.paragraph_format.line_spacing = Pt(self.config['title_line_spacing'])
                        
                        self._reset_pagination_properties(para_title)
                        self._format_heading(para_title, 1)
                
                # 格式化附件的副标题
                if att_subtitle_indices:
                    self._log(f"  > 识别到附件副标题（共 {len(att_subtitle_indices)} 行）")
                    for idx in att_subtitle_indices:
                        para_subtitle = all_blocks[idx]
                        self._log(f"    段落 {idx + 1}: 附件副标题行 - \"{para_subtitle.text.strip()[:30]}...\"")
                        self._strip_leading_whitespace(para_subtitle)
                        self._apply_font_to_runs(para_subtitle, self.config['subtitle_font'], self.config['subtitle_size'], set_color=apply_color)
                        para_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para_subtitle.paragraph_format.first_line_indent = None
                        
                        # 设置副标题行间距
                        spacing = para_subtitle._p.get_or_add_pPr().get_or_add_spacing()
                        spacing.set(qn('w:beforeAutospacing'), '0')
                        spacing.set(qn('w:afterAutospacing'), '0')
                        para_subtitle.paragraph_format.space_before = Pt(0)
                        para_subtitle.paragraph_format.space_after = Pt(0)
                        para_subtitle.paragraph_format.line_spacing = Pt(self.config['subtitle_line_spacing'])
                        
                        self._reset_pagination_properties(para_subtitle)
                
                # 计算下一个要处理的块索引
                if att_subtitle_indices:
                    next_idx = max(att_subtitle_indices) + 1
                elif att_title_indices:
                    next_idx = max(att_title_indices) + 1
                else:
                    next_idx = search_idx
                
                block_idx = next_idx
                continue
            
            elif re_h1.match(text_to_check):
                self._log(f"段落 {current_block_num}: 一级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._format_heading(para, 1)
                self._apply_font_to_runs(para, self.config['h1_font'], self.config['h1_size'], set_color=apply_color)
                self._apply_text_indent_and_align(para)
                self._reset_pagination_properties(para)

            elif re_h2.match(text_to_check):
                self._log(f"段落 {current_block_num}: 二级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                
                parts = para.text.split('。', 1)
                
                if len(parts) == 2 and parts[1].strip():
                    self._log("  > 检测到二级标题与正文在同一段落，执行段内格式拆分。")
                    title_len = len(parts[0]) + 1
                    
                    original_runs = []
                    for r in para.runs:
                        original_runs.append({
                            'text': r.text, 'bold': r.bold, 'italic': r.italic,
                            'underline': r.underline, 'font_color': r.font.color.rgb
                        })
                    
                    para.clear()

                    char_count = 0
                    for run_info in original_runs:
                        run_text = run_info['text']
                        run_end_pos = char_count + len(run_text)
                        
                        title_run, body_run, new_run = None, None, None

                        if run_end_pos <= title_len:
                            new_run = para.add_run(run_text)
                            self._set_run_font(new_run, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                        
                        elif char_count >= title_len:
                            new_run = para.add_run(run_text)
                            self._set_run_font(new_run, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                        
                        else:
                            split_index = title_len - char_count
                            title_part = run_text[:split_index]
                            body_part = run_text[split_index:]
                            
                            if title_part:
                                title_run = para.add_run(title_part)
                                self._set_run_font(title_run, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                            if body_part:
                                body_run = para.add_run(body_part)
                                self._set_run_font(body_run, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                        
                        runs_to_format = [r for r in [title_run, body_run] if r] or ([new_run] if new_run else [])
                        for r in runs_to_format:
                            if r:
                                r.bold = run_info['bold']; r.italic = run_info['italic']
                                r.underline = run_info['underline']
                                if run_info['font_color']: r.font.color.rgb = run_info['font_color']
                        
                        char_count = run_end_pos
                    
                    self._format_heading(para, 2)
                    self._apply_text_indent_and_align(para)
                    self._reset_pagination_properties(para)

                else:
                    match = re.match(r'^[（\(](.+?)[）\)](.*)', text_to_check, re.DOTALL)
                    if match and not (text_to_check.startswith('（') and text_to_check.strip().endswith('）')):
                        self._log("  > 已将二级标题的括号统一为中文括号。")
                        for r in para.runs: r.text = r.text.replace('(', '（', 1).replace(')', '）', 1)
                    self._format_heading(para, 2)
                    self._apply_font_to_runs(para, self.config['h2_font'], self.config['h2_size'], set_color=apply_color)
                    self._apply_text_indent_and_align(para)
                    self._reset_pagination_properties(para)
                    
            elif re_h3.match(text_to_check):
                self._log(f"段落 {current_block_num}: 三级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._format_heading(para, 3)
                self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                self._apply_text_indent_and_align(para)
                self._reset_pagination_properties(para)
                
            elif re_h4.match(text_to_check):
                self._log(f"段落 {current_block_num}: 四级标题 - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._format_heading(para, 4)
                self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                self._apply_text_indent_and_align(para)
                self._reset_pagination_properties(para)
                
            elif not is_from_txt:
                if para.alignment in [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]:
                    align_text = "居中" if para.alignment == WD_ALIGN_PARAGRAPH.CENTER else "右对齐"
                    self._log(f"段落 {current_block_num}: {align_text}正文 - 保留原对齐")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                    self._reset_pagination_properties(para)
                elif leading_space_count > 5:
                    self._log(f"段落 {current_block_num}: 正文 (保留前导空格) - \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    self._reset_pagination_properties(para)
                elif (para.paragraph_format.first_line_indent is None or para.paragraph_format.first_line_indent.pt == 0) and leading_space_count == 0:
                    self._log(f"段落 {current_block_num}: 正文 (保留0缩进) - \"{para_text_preview}...\"")
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    self._reset_pagination_properties(para)
                else:
                    self._log(f"段落 {current_block_num}: 正文 (应用标准缩进) - \"{para_text_preview}...\"")
                    self._strip_leading_whitespace(para)
                    self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                    self._apply_text_indent_and_align(para)
                    self._reset_pagination_properties(para)
            else:
                self._log(f"段落 {current_block_num}: 正文 (源自TXT，强制缩进) - \"{para_text_preview}...\"")
                self._strip_leading_whitespace(para)
                self._apply_font_to_runs(para, self.config['body_font'], self.config['body_size'], set_color=apply_color)
                self._apply_text_indent_and_align(para)
                self._reset_pagination_properties(para)
            
            block_idx += 1
        
        self._format_tables(doc, apply_color=apply_color)
        self._apply_page_setup(doc, is_from_txt=is_from_txt)
        self._log("正在保存最终文档...")
        doc.save(output_path)
